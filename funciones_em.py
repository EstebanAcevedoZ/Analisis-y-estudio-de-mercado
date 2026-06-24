# ==========================================
# Archivo: funciones_em.py
# Descripción: Motor de funciones para la generación de reportes de mercado.
# Autor: Esteban Acevedo Z.
# Fecha: 06/2026
# ==========================================

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import os
import re
import numpy as np
import requests
import geopandas as gpd
import contextily as ctx
from shapely.geometry import shape
import matplotlib
matplotlib.use('Agg')

LAYER_URL_LOTE = "https://serviciosgis.catastrobogota.gov.co/arcgis/rest/services/catastro/lote/MapServer/0"
VALOR_REFERENCIA_URL = "https://serviciosgis.catastrobogota.gov.co/arcgis/rest/services/catastro/valorreferencia/MapServer/0"
ALTURA_MEDIA_URL = "https://serviciosgis.catastrobogota.gov.co/arcgis/rest/services/catastro/alturamedia/MapServer/0"
CRS_WGS84 = "EPSG:4326"
CRS_METROS = "EPSG:3116"
IMAGES_OUTPUT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Imágenes")

MES_ESP_ABBR = {
    1: 'ene', 2: 'feb', 3: 'mar', 4: 'abr', 5: 'may', 6: 'jun',
    7: 'jul', 8: 'ago', 9: 'sep', 10: 'oct', 11: 'nov', 12: 'dic'
}

MES_ESP_FULL = {
    'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
    'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12,
    'ene': 1, 'feb': 2, 'mar': 3, 'abr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'ago': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dic': 12,
    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
}

def _parse_db_month_year(path):
    if not path:
        return None
    nombre = os.path.basename(str(path)).lower()
    match = re.search(r'\b(0[1-9]|1[0-2])-(\d{2})\b', nombre)
    if match:
        mes = int(match.group(1))
        ano = 2000 + int(match.group(2))
        return pd.Timestamp(year=ano, month=mes, day=1)

    match = re.search(r'\b(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre|ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\s*[\-_ ]*\s*(\d{4})\b', nombre)
    if match:
        mes = MES_ESP_FULL.get(match.group(1), None)
        ano = int(match.group(2))
        if mes:
            return pd.Timestamp(year=ano, month=mes, day=1)
    return None

def _find_column_by_patterns(columns, patterns):
    for pat in patterns:
        prog = re.compile(pat, re.IGNORECASE)
        for col in columns:
            if prog.search(str(col)):
                return col
    return None

def _rename_column_if_exists(df, new_name, patterns):
    col = _find_column_by_patterns(df.columns, patterns)
    if col and col != new_name:
        df.rename(columns={col: new_name}, inplace=True)
    return col

def _select_excel_engine(path):
    if isinstance(path, tuple):
        path = path[0]
    ext = os.path.splitext(str(path).lower())[1]
    if ext == '.xlsb':
        return 'pyxlsb'
    return None

def _rename_coordinate_columns(df, target_lat='Latitud', target_lon='Longitud'):
    existing_lower = {str(c).lower() for c in df.columns}
    rename_map = {}
    for col in df.columns:
        col_lower = str(col).lower()
        if 'latitud' in col_lower and target_lat.lower() not in existing_lower:
            rename_map[col] = target_lat
        if 'longitud' in col_lower and target_lon.lower() not in existing_lower:
            rename_map[col] = target_lon
    if rename_map:
        df.rename(columns=rename_map, inplace=True)

def _rename_latest_regex(df, patron_regex, new_name):
    meses_num = {'ene': 1, 'feb': 2, 'mar': 3, 'abr': 4, 'may': 5, 'jun': 6, 
                 'jul': 7, 'ago': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dic': 12}
    patron = re.compile(patron_regex, re.IGNORECASE)
    matches = []
    
    for col in df.columns:
        match = patron.search(str(col))
        if match:
            mes = meses_num[match.group(1).lower()]
            ano = 2000 + int(match.group(2))
            fecha = pd.Timestamp(year=ano, month=mes, day=1)
            matches.append((col, fecha))
            
    if matches:
        matches.sort(key=lambda x: x[1], reverse=True) 
        latest_col = matches[0][0]
        if latest_col != new_name:
            df.rename(columns={latest_col: new_name}, inplace=True)
            #print(f"[ÉXITO] Columna más reciente detectada: Se renombró '{latest_col}' a '{new_name}'")
        return new_name
    return None

def _rename_latest_cu(df, prefijo, new_name):
    patron = re.compile(rf'^\s*{re.escape(prefijo)}\s*(\d{{4}})-(\d{{2}})-(\d{{2}})\s*$', re.IGNORECASE)
    matches = []
    
    for col in df.columns:
        match = patron.search(str(col))
        if match:
            fecha = pd.to_datetime(match.group(0).replace(prefijo, '').strip(), errors='coerce')
            if pd.notna(fecha):
                matches.append((col, fecha))
                
    if matches:
        matches.sort(key=lambda x: x[1], reverse=True)
        latest_col = matches[0][0]
        if latest_col != new_name:
            df.rename(columns={latest_col: new_name}, inplace=True)
            #print(f"[ÉXITO] Columna más reciente detectada: Se renombró '{latest_col}' a '{new_name}'")
        return new_name
    return None

def query_geojson(layer_url, where=None, geometry=None):
    params = {
        "f": "geojson",
        "outFields": "*",
        "returnGeometry": "true",
        "where": where or "1=1"
    }

    if geometry:
        params.update({
            "geometry": geometry,
            "geometryType": "esriGeometryEnvelope",
            "inSR": 4326,
            "spatialRel": "esriSpatialRelIntersects"
        })

    r = requests.get(f"{layer_url}/query", params=params, timeout=60)
    r.raise_for_status()

    rows = []
    for f in r.json().get("features", []):
        props = f.get("properties", {})
        geom = f.get("geometry")
        if geom:
            try:
                props["geometry"] = shape(geom)
            except Exception:
                pass
        rows.append(props)

    if not rows:
        return gpd.GeoDataFrame(rows, crs=CRS_WGS84)

    df_rows = pd.DataFrame(rows)
    if "geometry" in df_rows.columns:
        return gpd.GeoDataFrame(df_rows, geometry="geometry", crs=CRS_WGS84)
    else:
        return gpd.GeoDataFrame(df_rows, crs=CRS_WGS84)

def generar_mapa_lote(doc, lotcodigo, escala=2000, dpi=300, ancho_cm=16, output_folder=None):
    output_folder = output_folder or IMAGES_OUTPUT_FOLDER
    os.makedirs(output_folder, exist_ok=True)

    gdf = query_geojson(LAYER_URL_LOTE, where=f"LOTCODIGO = '{lotcodigo}'")
    if gdf.empty:
        raise ValueError(f"No se encontró el lote {lotcodigo}")

    gdf = gdf.iloc[[0]].copy()
    gdf_m = gdf.to_crs(CRS_METROS)
    centroide = gdf_m.geometry.iloc[0].centroid

    ancho_mapa_m = (ancho_cm / 100) * escala
    mitad_ventana = ancho_mapa_m / 2
    fig_size_in = ancho_cm / 2.54

    fig, ax = plt.subplots(figsize=(fig_size_in, fig_size_in), frameon=False)
    ax.set_position([0, 0, 1, 1])

    gdf_3857 = gdf_m.to_crs(epsg=3857)
    centroide_3857 = gdf_3857.geometry.iloc[0].centroid

    ax.set_xlim(centroide_3857.x - mitad_ventana, centroide_3857.x + mitad_ventana)
    ax.set_ylim(centroide_3857.y - mitad_ventana, centroide_3857.y + mitad_ventana)

    ctx.add_basemap(ax, source=ctx.providers.CartoDB.Voyager)

    gdf_3857.plot(ax=ax, facecolor="none", edgecolor="red", linewidth=2.5)
    ax.scatter(centroide_3857.x, centroide_3857.y, s=35, color="red", edgecolors="white", linewidth=0.8, zorder=10)

    ax.annotate(
        "N",
        xy=(0.92, 0.88),
        xytext=(0.92, 0.78),
        xycoords="axes fraction",
        ha="center",
        va="center",
        fontsize=13,
        fontweight="bold",
        arrowprops=dict(facecolor="black", edgecolor="black", width=4, headwidth=12, headlength=12)
    )

    ax.text(
        0.03,
        0.03,
        "Escala 1:4.000",
        transform=ax.transAxes,
        fontsize=9,
        bbox=dict(facecolor="white", alpha=0.85, edgecolor="black", linewidth=0.5)
    )

    ax.set_axis_off()
    ax.set_aspect("equal")
    plt.subplots_adjust(left=0, right=1, top=1, bottom=0)

    output_png = os.path.join(output_folder, f"mapa_lote_{lotcodigo}.png")
    plt.savefig(output_png, dpi=dpi, bbox_inches="tight", pad_inches=0)
    plt.close()
    doc.add_picture(output_png, width=Inches(6.0))
    return output_png

def generar_mapa_manzanas(doc,layer_url, mancodigo, campo_valor, nombre_salida, distancia_m=200, ano=2026, output_folder=None, dpi=300, alpha=0.7):
    output_folder = output_folder or IMAGES_OUTPUT_FOLDER
    os.makedirs(output_folder, exist_ok=True)

    output_png = os.path.join(output_folder, f"{nombre_salida}_{mancodigo}_{ano}.png")
    output_csv = os.path.join(output_folder, f"{nombre_salida}_{mancodigo}_{ano}.csv")

    filtro_ano = f"ANO = DATE '{ano}-01-01'"
    gdf_target = query_geojson(layer_url, where=f"MANCODIGO = '{mancodigo}' AND {filtro_ano}")
    if gdf_target.empty:
        raise ValueError(f"No se encontró ninguna manzana con MANCODIGO = {mancodigo}")

    gdf_target = gdf_target.iloc[[0]].copy()
    target_m = gdf_target.to_crs(CRS_METROS)
    centroide_m = target_m.geometry.iloc[0].centroid
    buffer_m = centroide_m.buffer(distancia_m)

    gdf_buffer = gpd.GeoDataFrame([{"geometry": buffer_m}], geometry="geometry", crs=CRS_METROS)
    buffer_wgs84 = gdf_buffer.to_crs(CRS_WGS84)
    minx, miny, maxx, maxy = buffer_wgs84.total_bounds
    envelope = f"{minx},{miny},{maxx},{maxy}"

    gdf_candidates = query_geojson(layer_url, where=filtro_ano, geometry=envelope)
    if gdf_candidates.empty:
        raise ValueError("No se encontraron manzanas candidatas alrededor del buffer.")

    candidates_m = gdf_candidates.to_crs(CRS_METROS)
    gdf_selected_m = candidates_m[candidates_m.intersects(buffer_m)].copy()
    if gdf_selected_m.empty:
        raise ValueError("No se encontraron manzanas seleccionadas.")

    gdf_selected_m.drop(columns="geometry").to_csv(output_csv, index=False, encoding="utf-8-sig")
    gdf_selected_3857 = gdf_selected_m.to_crs(epsg=3857)
    gdf_target_3857 = gdf_target.to_crs(epsg=3857)
    gdf_buffer_3857 = gdf_buffer.to_crs(epsg=3857)

    cmap = "viridis"
    if campo_valor == "ALTURA":
        cmap = "YlOrBr"

    fig, ax = plt.subplots(figsize=(10, 10))
    gdf_selected_3857.plot(ax=ax, column=campo_valor if campo_valor in gdf_selected_3857.columns else None, cmap=cmap, edgecolor="black", linewidth=0.25, alpha=alpha, legend=False)

    for _, row in gdf_selected_3857.iterrows():
        if row.geometry is None or row.geometry.is_empty:
            continue
        valor = row.get(campo_valor)
        if valor is None:
            continue
        punto = row.geometry.representative_point()
        etiqueta = f"${valor:,.0f}" if campo_valor == "V_REF" else f"{int(round(valor))}"
        ax.text(punto.x, punto.y, etiqueta, fontsize=11, color="black", ha="center", va="center", zorder=20, fontweight="bold")

    gdf_buffer_3857.boundary.plot(ax=ax, linewidth=1.5, linestyle="--", color="red")
    gdf_target_3857.plot(ax=ax, facecolor="none", edgecolor="red", linewidth=2.5)
    ctx.add_basemap(ax, source=ctx.providers.CartoDB.Voyager)
    ax.set_axis_off()
    plt.tight_layout()
    plt.savefig(output_png, dpi=dpi, bbox_inches="tight", pad_inches=0.05)
    plt.close()
    doc.add_picture(output_png, width=Inches(6.0))
    return output_png, output_csv, len(gdf_selected_m)

def generar_mapa_valor_referencia(doc, mancodigo, distancia_m=200, ano=2026, output_folder=None, dpi=300, alpha=0.7):
    valor_png, _, _ = generar_mapa_manzanas(doc, VALOR_REFERENCIA_URL, mancodigo, "V_REF", "mapa_valor_referencia", distancia_m, ano, output_folder, dpi, alpha)
    return valor_png

def generar_mapa_altura_media(doc, mancodigo, distancia_m=200, ano=2026, output_folder=None, dpi=300, alpha=0.7):
    altura_png, _, _ = generar_mapa_manzanas(doc, ALTURA_MEDIA_URL, mancodigo, "ALTURA", "mapa_altura_media", distancia_m, ano, output_folder, dpi, alpha)
    return altura_png

def crear_precio_por_m2(df, col_area='Area', prefijo_precio='P'):
    df = df.copy()
    if col_area not in df.columns:
        print(f"Error: La columna '{col_area}' no existe en el DataFrame")
        return df
    
    patron = re.compile(rf'^{re.escape(prefijo_precio)}(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\d{{2}}$', re.IGNORECASE)
    cols_precio = [c for c in df.columns if patron.match(str(c))]
    
    if not cols_precio:
        print(f"Advertencia: No se encontraron columnas de precio con prefijo '{prefijo_precio}'")
        return df
    
    nuevas_cols = {}
    for col_precio in cols_precio:
        sufijo = col_precio[len(prefijo_precio):]
        nombre_nueva_col = f"Pm{sufijo.lower()}"
        nuevas_cols[nombre_nueva_col] = (df[col_precio] / df[col_area]).replace([np.inf, -np.inf], np.nan)
    
    if nuevas_cols:
        df = pd.concat([df, pd.DataFrame(nuevas_cols)], axis=1)
    
    return df

def _extraer_fecha_de_columna(col, formato):
    meses_num = {'ene': 1, 'feb': 2, 'mar': 3, 'abr': 4, 'may': 5, 'jun': 6, 
                 'jul': 7, 'ago': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dic': 12}
    
    if formato == 'texto_fecha':
        fecha_limpia = str(col).lower().replace('00:00:00', '').strip()
        match = re.search(r'(\d{4}-\d{2}-\d{2})', fecha_limpia)
        if match:
            fecha_str = match.group(1)
            return pd.to_datetime(fecha_str, errors='coerce')
        return pd.NaT
        
    elif formato == 'prefijo_mes':
        c_str = str(col).lower()
        match = re.search(r'(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})', c_str)
        if match:
            mes_str = match.group(1)
            año_str = match.group(2)
            if mes_str in meses_num:
                año = 2000 + int(año_str)
                mes = meses_num[mes_str]
                return pd.to_datetime(f"{año}-{mes:02d}-01")
        return pd.NaT
    
    return pd.NaT

def grafico_evolucion_temporal(
    doc, 
    dict_dfs, 
    lista_ids, 
    cols_agrupacion, 
    nombre_base='Base Formulada', 
    col_id='Codproyecto',
    tipo_columnas='vendidas',
    prefijo_columnas=None,
    formato_columnas=None,
    etiqueta_y='Unidades',
    etiqueta_tabla_total='Total (últ. 12 meses)',
    etiqueta_tabla_promedio='Promedio mensual (Últ. 12 meses)',
    incluir_tabla=True
):
    if tipo_columnas == 'vendidas':
        prefijo_columnas = prefijo_columnas or 'vendidas'
        formato_columnas = formato_columnas or 'texto_fecha'
    elif tipo_columnas == 'v_mes':
        prefijo_columnas = prefijo_columnas or 'V'
        formato_columnas = formato_columnas or 'prefijo_mes'
    elif tipo_columnas == 'custom':
        if prefijo_columnas is None:
            print("Error: Para tipo_columnas='custom', debe especificar prefijo_columnas")
            return
        formato_columnas = formato_columnas or 'texto_fecha'
    else:
        print(f"Error: tipo_columnas='{tipo_columnas}' no es válido.")
        return
    
    df = dict_dfs.get(nombre_base)
    if df is None:
        print(f"Error: Base de datos '{nombre_base}' no encontrada en diccionario.")
        return
    
    lista_ids_str = [str(x) for x in lista_ids]
    df_filtrado = df[df[col_id].astype(str).isin(lista_ids_str)]
    
    if df_filtrado.empty:
        matched_ids = set(df[df[col_id].astype(str).isin(lista_ids_str)][col_id].astype(str).unique())
        unmatched = [x for x in lista_ids_str if x not in matched_ids][:10]
        print(f"Error: No hay proyectos en esta zona para graficar. nombre_base={nombre_base}, col_id={col_id}")
        return

    if formato_columnas == 'prefijo_mes':
        patron = re.compile(rf'^{re.escape(prefijo_columnas)}(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\d{{2}}$', re.IGNORECASE)
        cols_temporal = [c for c in df_filtrado.columns if patron.match(str(c))]
    else:
        cols_temporal = [c for c in df_filtrado.columns if prefijo_columnas.lower() in str(c).lower()]

    if not cols_temporal:
        print(f"Advertencia: No se encontraron columnas con prefijo '{prefijo_columnas}' en formato '{formato_columnas}'")
        return

    fechas_validas = []
    for col in cols_temporal:
        fecha_obj = _extraer_fecha_de_columna(col, formato_columnas)
        if pd.notna(fecha_obj):
            fechas_validas.append((col, fecha_obj))

    if not fechas_validas:
        print("Error: No se pudieron parsear las fechas de las columnas.")
        return

    fechas_validas.sort(key=lambda x: x[1])
    ultimos_12_tuplas = fechas_validas[-12:]
    ultimas_12_cols = [x[0] for x in ultimos_12_tuplas]

    es_metrica_promedio = any(p in prefijo_columnas.lower() for p in ['pm', 'precio', 'area', 'm2', 'alc', 'baño'])

    if es_metrica_promedio:
        df_agrupado = df_filtrado.groupby(cols_agrupacion)[ultimas_12_cols].mean().reset_index()
        df_agrupado['Promedio Mensual (12m)'] = df_agrupado[ultimas_12_cols].mean(axis=1)
    else:
        df_agrupado = df_filtrado.groupby(cols_agrupacion)[ultimas_12_cols].sum().reset_index()
        df_agrupado['Promedio Mensual (12m)'] = df_agrupado[ultimas_12_cols].sum(axis=1) / 12

    fig, ax = plt.subplots(figsize=(8, 4.5))
    meses_es = {1:'Ene', 2:'Feb', 3:'Mar', 4:'Abr', 5:'May', 6:'Jun', 
                7:'Jul', 8:'Ago', 9:'Sep', 10:'Oct', 11:'Nov', 12:'Dic'}
    
    etiquetas_x = [f"{meses_es[fecha.month]}-{fecha.strftime('%y')}" for _, fecha in ultimos_12_tuplas]
    colors_palette = ["#ddff86", "#594eea","#255b46"]
    
    for idx, (_, row) in enumerate(df_agrupado.iterrows()):
        if len(cols_agrupacion) == 1:
            label = str(row[cols_agrupacion[0]])
        else:
            label = " - ".join([str(row[c]) for c in cols_agrupacion])
        
        valores_y = row[ultimas_12_cols].values
        color = colors_palette[idx % len(colors_palette)]
        ax.plot(etiquetas_x, valores_y, marker='o', linewidth=2.5, markersize=6, label=label, color=color)

    ax.set_ylabel(etiqueta_y)
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.legend(title=" & ".join(cols_agrupacion))
    
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile, format='png', dpi=200, bbox_inches='tight')
    plt.close()
    memfile.seek(0)
    doc.add_picture(memfile, width=Inches(6.0))

    if incluir_tabla:
        doc.add_paragraph(f"Resumen de datos y promedio:")
        table = doc.add_table(rows=1, cols=len(cols_agrupacion) + 2)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        for i, col in enumerate(cols_agrupacion):
            hdr[i].text = col
        hdr[-2].text = "Promedio del año" if es_metrica_promedio else etiqueta_tabla_total
        hdr[-1].text = "Mediana del año" if es_metrica_promedio else etiqueta_tabla_promedio

        for _, row in df_agrupado.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(cols_agrupacion):
                row_cells[i].text = limpiar_texto(row[col])
                
            if es_metrica_promedio:
                val_total = row[ultimas_12_cols].mean()
                val_promedio = row[ultimas_12_cols].median()
            else:
                val_total = row[ultimas_12_cols].sum()
                val_promedio = row['Promedio Mensual (12m)']
            
            row_cells[-2].text = f"{val_total:,.1f}"
            row_cells[-1].text = f"{val_promedio:,.1f}"

def grafico_historico_doble_eje(
    doc,
    dict_dfs,
    lista_ids,
    cols_agrupacion,
    nombre_base='Base Formulada',
    col_id='Codproyecto',
    prefijo_barras='V',
    formato_barras='prefijo_mes',
    prefijo_linea='D',
    formato_linea='prefijo_mes',
    etiqueta_y_barras='Ventas Totales (unidades)',
    etiqueta_y_linea='Oferta Disponible (unidades)',
    titulo_barras='Ventas Totales',
    titulo_linea='Oferta Disponible',
    titulo_grafico='Evolución Histórica'
):
    df = dict_dfs.get(nombre_base)
    if df is None:
        print(f"Error: Base '{nombre_base}' no encontrada.")
        return

    lista_ids_str = [str(x) for x in lista_ids]
    df_filtrado = df[df[col_id].astype(str).isin(lista_ids_str)].copy()

    if df_filtrado.empty:
        print(f"Error: No hay proyectos para graficar doble eje en {nombre_base}")
        return

    def _encontrar_cols_y_fechas(prefix, formato):
        if formato == 'prefijo_mes':
            patron = re.compile(rf'^{re.escape(prefix)}(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\d{{2}}$', re.IGNORECASE)
            cols = [c for c in df_filtrado.columns if patron.match(str(c))]
        else:
            cols = [c for c in df_filtrado.columns if prefix.lower() in str(c).lower()]
        
        fechas = []
        for col in cols:
            f_obj = _extraer_fecha_de_columna(col, formato)
            if pd.notna(f_obj):
                fechas.append((col, f_obj))
        return fechas

    fechas_barras = _encontrar_cols_y_fechas(prefijo_barras, formato_barras)
    fechas_linea = _encontrar_cols_y_fechas(prefijo_linea, formato_linea)

    dict_b = {f: col for col, f in fechas_barras}
    dict_l = {f: col for col, f in fechas_linea}
    fechas_comunes = sorted(list(set(dict_b.keys()) & set(dict_l.keys())))

    if not fechas_comunes:
        print(f"Advertencia: No se pudieron alinear fechas para doble eje en {nombre_base}")
        return

    ultimos_12_fechas = fechas_comunes[-12:]
    cols_barras = [dict_b[f] for f in ultimos_12_fechas]
    cols_linea = [dict_l[f] for f in ultimos_12_fechas]

    df_agrup_b = df_filtrado.groupby(cols_agrupacion)[cols_barras].sum()
    df_agrup_l = df_filtrado.groupby(cols_agrupacion)[cols_linea].sum()

    meses_es = {1:'Ene', 2:'Feb', 3:'Mar', 4:'Abr', 5:'May', 6:'Jun', 
                7:'Jul', 8:'Ago', 9:'Sep', 10:'Oct', 11:'Nov', 12:'Dic'}
    etiquetas_x = [f"{meses_es[f.month]}-{f.strftime('%y')}" for f in ultimos_12_fechas]
    
    for idx_cat, index_val in enumerate(df_agrup_b.index):
        fig, ax1 = plt.subplots(figsize=(9, 5))
        
        categoria_label = str(index_val) if isinstance(index_val, (str, int)) else " - ".join(map(str, index_val))
        
        valores_b = df_agrup_b.loc[index_val].values
        valores_l = df_agrup_l.loc[index_val].values
        
        color_b = '#ddff86' 
        color_l = '#594eea' 

        bars = ax1.bar(etiquetas_x, valores_b, color=color_b, alpha=0.85, width=0.55, label=titulo_barras)
        ax1.set_ylabel(etiqueta_y_barras, fontweight='bold', fontsize=10)
        ax1.tick_params(axis='y')
        ax1.grid(True, linestyle=':', alpha=0.6, axis='y')

        ax2 = ax1.twinx()
        line = ax2.plot(etiquetas_x, valores_l, color=color_l, marker='o', linewidth=2.5, markersize=7, label=titulo_linea)
        ax2.set_ylabel(etiqueta_y_linea, fontweight='bold', fontsize=10)
        ax2.tick_params(axis='y')

        for bar in bars:
            height = bar.get_height()
            ax1.annotate(f'{int(round(height))}',
                        xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 4), 
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=9, fontweight='bold')

        for i, val in enumerate(valores_l):
            ax2.annotate(f'{int(round(val))}',
                        xy=(etiquetas_x[i], valores_l[i]),
                        xytext=(0, 8), 
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=9, color=color_l, fontweight='bold')

        ax1.spines['top'].set_visible(False)
        ax2.spines['top'].set_visible(False)

        handles1, labels1 = ax1.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(handles1 + handles2, labels1 + labels2, loc='upper left', framealpha=0.9)

        ax1.set_title(f"{titulo_grafico} - {categoria_label}", fontsize=12, fontweight='bold', pad=15)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()

        memfile = BytesIO()
        plt.savefig(memfile, format='png', dpi=200, bbox_inches='tight')
        plt.close()
        memfile.seek(0)
        doc.add_picture(memfile, width=Inches(6.0))
        doc.add_paragraph()

def tabla_proporciones(doc, dict_dfs, lista_ids, config_campos, col_id='Codproyecto'):
    if not lista_ids:
        doc.add_paragraph("No hay proyectos en la zona para calcular proporciones.")
        return

    resultados = []
    for cfg in config_campos:
        if len(cfg) == 4:
            titulo_word, nombre_base, col_excel, es_dummy = cfg
        else:
            titulo_word, nombre_base, col_excel = cfg
            es_dummy = False
            
        df = dict_dfs.get(nombre_base)
        
        if df is None or col_excel not in df.columns:
            continue
            
        df_filtrado = df[df[col_id].isin(lista_ids)]
        serie_datos = df_filtrado[col_excel].dropna()
        
        if not serie_datos.empty:
            if es_dummy:
                mapeo_dummies = {
                    1: 'Sí', 1.0: 'Sí', '1': 'Sí', '1.0': 'Sí', True: 'Sí', 'Si': 'Sí', 'SI': 'Sí',
                    0: 'No', 0.0: 'No', '0': 'No', '0.0': 'No', False: 'No', 'NO': 'No'
                }
                serie_datos = serie_datos.map(lambda x: mapeo_dummies.get(x, str(x)))
                
            proporciones = serie_datos.value_counts(normalize=True) * 100
            resultados.append((titulo_word, proporciones))
            
    if not resultados:
        doc.add_paragraph("No se encontró información disponible para las variables solicitadas.")
        return
        
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Detalle"
    hdr[2].text = "Participación"
    
    for titulo_word, proporciones in resultados:
        es_primera_fila = True
        
        for valor, pct in proporciones.items():
            row_cells = table.add_row().cells
            
            if es_primera_fila:
                p = row_cells[0].paragraphs[0]
                p.add_run(titulo_word).bold = True
                es_primera_fila = False
            row_cells[1].text = limpiar_texto(valor)
            row_cells[2].text = f"{pct:.1f}%"

def obtener_proyectos_cercanos(df_proyectos, coordenada, radio_km, filtros="default", min_available_ratio=None, min_available_units=None, df_inmuebles=None, df_inmuebles_proj_id='Codproyecto', df_inmuebles_col_available='Dactual', df_inmuebles_col_total='Oferta Total Inm.', fila_cod='Codproyecto', excluir_codigos=None):
    df_proyectos = df_proyectos.copy()
    
    if filtros == "default":
        filtros = [('Activo', '==', 'Si')]
    elif filtros is None:
        filtros = []

    lat1_rad = np.radians(coordenada[0])
    lon1_rad = np.radians(coordenada[1])
    lat2_rad = np.radians(df_proyectos['Latitud'])
    lon2_rad = np.radians(df_proyectos['Longitud'])
    
    dlat = lat2_rad - lat1_rad
    dlon = lon2_rad - lon1_rad
    
    a = np.sin(dlat/2)**2 + np.cos(lat1_rad) * np.cos(lat2_rad) * np.sin(dlon/2)**2
    c = 2 * np.arcsin(np.sqrt(a))
    
    df_proyectos['distancia'] = 6371.0 * c
    mascara = df_proyectos['distancia'] <= radio_km
    
    for col, operador, valor in filtros:
        if col in df_proyectos.columns:
            col_data = df_proyectos[col]
            es_valor_numerico = isinstance(valor, (int, float)) and not isinstance(valor, bool)
            
            if es_valor_numerico:
                col_data_procesada = pd.to_numeric(col_data, errors='coerce')
                valor_comparar = valor
            else:
                col_data_procesada = col_data.astype(str).str.strip()
                valor_comparar = str(valor).strip()
            
            if operador == '==': 
                mascara &= (col_data_procesada == valor_comparar)
            elif operador == '!=': 
                mascara &= (col_data_procesada != valor_comparar) & (col_data_procesada.notna())
            elif operador == '>': 
                mascara &= (col_data_procesada > valor_comparar)
            elif operador == '<': 
                mascara &= (col_data_procesada < valor_comparar)
            elif operador == '>=': 
                mascara &= (col_data_procesada >= valor_comparar)
            elif operador == '<=': 
                mascara &= (col_data_procesada <= valor_comparar)
        else:
            print(f"Advertencia: Columna '{col}' no encontrada en DataFrame. Columnas disponibles: {list(df_proyectos.columns[:10])}")
    
    df_filtrado = df_proyectos[mascara]
    if fila_cod not in df_filtrado.columns:
        print(f"Error: No se encontró la columna '{fila_cod}'.")
        return []
    
    ids_preliminares = df_filtrado[fila_cod].unique().tolist()
    
    if excluir_codigos is not None and len(excluir_codigos) > 0:
        print(f"Intentando excluir los siguientes códigos: {excluir_codigos}")
        codigos_a_excluir_set = set()
        for cod in excluir_codigos:
            cod_str = str(cod).strip().lower()
            codigos_a_excluir_set.add(cod_str)
        
        ids_preliminares_originales = ids_preliminares.copy()
        ids_preliminares = []
        for pid in ids_preliminares_originales:
            pid_str = str(pid).strip().lower()
            if pid_str not in codigos_a_excluir_set:
                ids_preliminares.append(pid)

    if df_inmuebles is not None and (min_available_ratio is not None or min_available_units is not None):
        col_proy = df_inmuebles_proj_id
        col_available = df_inmuebles_col_available
        col_total = df_inmuebles_col_total

        if col_proy in df_inmuebles.columns and col_available in df_inmuebles.columns and col_total in df_inmuebles.columns:
            agrupado = df_inmuebles[df_inmuebles[col_proy].isin(ids_preliminares)].groupby(col_proy)[[col_available, col_total]].sum()
            oferta_segura = agrupado[col_total].replace(0, np.nan)
            agrupado['ratio'] = agrupado[col_available] / oferta_segura

            if min_available_ratio is not None and min_available_units is not None:
                ids_validos = agrupado[((agrupado['ratio'] >= min_available_ratio) & (agrupado[col_available] >= min_available_units)) | agrupado['ratio'].isna()].index.tolist()
            elif min_available_ratio is not None:
                ids_validos = agrupado[(agrupado['ratio'] >= min_available_ratio) | agrupado['ratio'].isna()].index.tolist()
            elif min_available_units is not None:
                ids_validos = agrupado[agrupado[col_available] >= min_available_units].index.tolist()
            else:
                ids_validos = ids_preliminares

            ids_preliminares = [pid for pid in ids_preliminares if pid in ids_validos]
        else:
            print(f"Advertencia: No se pudo calcular la disponibilidad. Faltan las columnas '{col_proy}', '{col_available}' o '{col_total}' en df_inmuebles.")

    return ids_preliminares

def separar_coordenadas(df, col_origen, col_lat='Latitud', col_lon='Longitud'):
    if col_origen not in df.columns:
        print(f"Error: La columna '{col_origen}' no existe en la base de datos.")
        return df

    split_coords = df[col_origen].astype(str).str.split(',', n=1, expand=True)

    if split_coords.shape[1] == 2:
        df[col_lat] = pd.to_numeric(split_coords[0].str.strip(), errors='coerce')
        df[col_lon] = pd.to_numeric(split_coords[1].str.strip(), errors='coerce')
    else:
        print(f"Advertencia: No se encontraron comas en '{col_origen}' para separar.")

    return df

def limpiar_texto(texto):
    """Elimina caracteres invisibles de control que corrompen el XML de Word."""
    if pd.isna(texto):
        return "N/D"
    texto_str = str(texto)
    # Regex para eliminar todos los caracteres ilegales en XML 1.0
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]', '', texto_str)

def formatear_valor(valor, es_dinero=False):
    """Añade formato de miles y limpia caracteres ilegales."""
    try:
        num = float(valor)
        if es_dinero:
            return f"${num:,.0f}".replace(",", ".")
        return f"{num:,.0f}".replace(",", ".")
    except (ValueError, TypeError):
        # Si no es un número (es texto), lo pasamos por el filtro limpiador
        return limpiar_texto(valor)

def _consolidar_bases(dict_dfs, lista_ids, config_columnas, cols_agrupacion, col_id='Codproyecto'):
    bases_a_cruzar = {}
    config_interna = {}
    
    for nombre_visual, cfg in config_columnas.items():
        if len(cfg) == 5:
            base, col, dinero, met, extra = cfg
        else:
            base = None
            col, dinero, met, extra = cfg 
            
        if met == 'calc':
            config_interna[nombre_visual] = (None, dinero, met, extra)
            continue
            
        if base in dict_dfs:
            if base not in bases_a_cruzar:
                bases_a_cruzar[base] = set()
            
            if col_id in dict_dfs[base].columns:
                bases_a_cruzar[base].add(col_id)
            if col: 
                bases_a_cruzar[base].add(col)
                if met == 'sum_last12_sells':
                    palabra_clave = str(col).split(' ')[0].lower()
                    for c_df in dict_dfs[base].columns:
                        if isinstance(c_df, str) and palabra_clave in c_df.lower():
                            if re.search(r'\d{4}-\d{2}-\d{2}', c_df):
                                bases_a_cruzar[base].add(c_df)

            if isinstance(extra, str): bases_a_cruzar[base].add(extra)
            
            col_rename = col if col in cols_agrupacion or col == col_id else f"{base}__{col}"
            extra_rename = extra if isinstance(extra, str) and (extra in cols_agrupacion or extra == col_id) else (f"{base}__{extra}" if isinstance(extra, str) else None)
            
            config_interna[nombre_visual] = (col_rename, dinero, met, extra_rename)
        else:
            config_interna[nombre_visual] = (col, dinero, met, extra)

    if not bases_a_cruzar:
        return pd.DataFrame(), config_interna

    for base in bases_a_cruzar:
        for c_grp in cols_agrupacion:
            if c_grp in dict_dfs[base].columns:
                bases_a_cruzar[base].add(c_grp)

    df_unificado = None
    for base, cols in bases_a_cruzar.items():
        df_b = dict_dfs[base]
        valid_cols = [c for c in cols if c in df_b.columns]
        
        if col_id in df_b.columns:
            df_sub = df_b[df_b[col_id].isin(lista_ids)][valid_cols].copy()
        else:
            df_sub = df_b[valid_cols].copy()
        
        renames = {}
        for c in valid_cols:
            if c != col_id and c not in cols_agrupacion:
                renames[c] = f"{base}__{c}"
        df_sub = df_sub.rename(columns=renames)
        
        if df_unificado is None:
            df_unificado = df_sub
        else:
            join_keys = [c for c in valid_cols if c == col_id or c in cols_agrupacion]
            join_keys = [c for c in join_keys if c in df_unificado.columns]
            
            if join_keys:
                df_unificado = pd.merge(df_unificado, df_sub, on=join_keys, how='outer')
            else:
                df_unificado = pd.concat([df_unificado, df_sub], axis=0, ignore_index=True)
            
    return df_unificado, config_interna

def tabla_proyectos(doc, dict_dfs, lista_ids, config_columnas, cols_agrupacion, col_id='Codproyecto'):
    df_filtrado, config_interna = _consolidar_bases(dict_dfs, lista_ids, config_columnas, cols_agrupacion, col_id)
    
    if df_filtrado.empty:
        doc.add_paragraph("No se encontraron datos para mostrar en la tabla.")
        return

    def calcular_grupo(grupo):
        resultado = {}
        for nombre_visual, (col_excel, es_dinero, metodo, col_peso) in config_interna.items():
            if metodo == 'calc':
                continue
                
            if col_excel in cols_agrupacion:
                if len(cols_agrupacion) == 1:
                    resultado[nombre_visual] = grupo.name
                else:
                    idx = cols_agrupacion.index(col_excel)
                    resultado[nombre_visual] = grupo.name[idx]
                continue
                
            special_methods = ['sum_last12_sells']
            if col_excel not in grupo.columns and col_excel is not None and metodo not in special_methods:
                resultado[nombre_visual] = np.nan
                continue
                
            if metodo == 'first':
                resultado[nombre_visual] = grupo[col_excel].iloc[0]
            elif metodo == 'sum':
                resultado[nombre_visual] = grupo[col_excel].sum(skipna=True)
            elif metodo == 'sum_last12_sells':
                if '__' in str(col_excel):
                    prefijo_base = str(col_excel).split('__')[0] + '__'
                else:
                    prefijo_base = ''
                    
                cols_coincidentes = []
                for c in grupo.columns:
                    c_str = str(c)
                    c_clean = c_str.split('__')[1] if '__' in c_str else c_str
                    c_clean_lower = c_clean.lower()
                    
                    if re.search(r'\d{4}-\d{2}-\d{2}', c_clean_lower):
                        if any(kw in c_clean_lower for kw in ['vend', 'vent']):
                            if not any(exc in c_clean_lower for exc in ['precio', 'metro', 'saldo', 'acum', 'total', 'promedio', 'vabr']):
                                fecha = pd.to_datetime(re.search(r'\d{4}-\d{2}-\d{2}', c_clean_lower).group(), errors='coerce')
                                if pd.notna(fecha):
                                    cols_coincidentes.append((c_str, fecha))
                                
                cols_coincidentes.sort(key=lambda x: x[1])
                ultimos_12 = [c[0] for c in cols_coincidentes[-12:]]
                
                if ultimos_12:
                    resultado[nombre_visual] = grupo[ultimos_12].sum(numeric_only=True).sum(skipna=True)
                else:
                    resultado[nombre_visual] = np.nan
            elif metodo == 'mean':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                # Si es dinero o área, el 0 es un dato faltante (NaN)
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.mean(skipna=True)
                
            elif metodo == 'median':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.median(skipna=True)
                
            elif metodo == 'max':
                resultado[nombre_visual] = pd.to_numeric(grupo[col_excel], errors='coerce').max(skipna=True)
                
            elif metodo == 'min':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.min(skipna=True)
                
            elif metodo == 'avg12':
                resultado[nombre_visual] = grupo[col_excel].sum(skipna=True) / 12
                
            elif metodo in ['weighted_mean', 'weighted_average']:
                if col_peso and col_peso in grupo.columns:
                    df_valido = grupo.copy()
                    
                    # 1. Convertir a numérico
                    df_valido[col_excel] = pd.to_numeric(df_valido[col_excel], errors='coerce')
                    df_valido[col_peso] = pd.to_numeric(df_valido[col_peso], errors='coerce')
                    
                    # 2. Convertir ceros a NaN si es dinero o área
                    if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower():
                        df_valido[col_excel] = df_valido[col_excel].replace(0, np.nan)
                    
                    # 3. Eliminar cualquier fila que ahora sea NaN
                    df_valido = df_valido.dropna(subset=[col_excel, col_peso])
                    
                    valores_num = df_valido[col_excel]
                    pesos_num = df_valido[col_peso]
                    
                    # 4. Ignorar filas cuyo "peso" (ej. Unidades Disponibles) sea 0 o negativo
                    mascara_valida = (pesos_num > 0)
                    
                    if mascara_valida.sum() > 0:
                        resultado[nombre_visual] = np.average(valores_num[mascara_valida], weights=pesos_num[mascara_valida])
                    else:
                        resultado[nombre_visual] = np.nan
                else:
                    vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                    if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                        vals = vals.replace(0, np.nan)
                    resultado[nombre_visual] = vals.mean(skipna=True)
            elif metodo == 'concat':
                unique_values = grupo[col_excel].dropna().unique()
                if len(unique_values) > 1:
                    resultado[nombre_visual] = ', '.join(map(str, unique_values))
                elif len(unique_values) == 1:
                    resultado[nombre_visual] = unique_values[0]
                else:
                    resultado[nombre_visual] = np.nan
            else:
                resultado[nombre_visual] = grupo[col_excel].iloc[0]
                
        for nombre_visual, (col_excel, es_dinero, metodo, operacion) in config_interna.items():
            if metodo == 'calc':
                try:
                    resultado[nombre_visual] = operacion(resultado)
                except (ZeroDivisionError, KeyError, TypeError):
                    resultado[nombre_visual] = np.nan      
        return pd.Series(resultado)
    
    try:
        df_agrupado = df_filtrado.groupby(cols_agrupacion).apply(calcular_grupo, include_groups=False)
    except TypeError:
        df_agrupado = df_filtrado.groupby(cols_agrupacion).apply(calcular_grupo)
        
    df_agrupado = df_agrupado.reset_index(drop=True)
    
    num_cols = len(config_columnas)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    nombres_columnas_keys = list(config_columnas.keys())
    
    for i, nombre_visual in enumerate(nombres_columnas_keys):
        hdr_cells[i].text = nombre_visual

    for _, row in df_agrupado.iterrows():
        row_cells = table.add_row().cells
        for i, nombre_visual in enumerate(nombres_columnas_keys):
            valor = row.get(nombre_visual, np.nan)
            es_dinero = config_columnas[nombre_visual][2] 
            
            if pd.isna(valor) or (isinstance(valor, (int, float)) and np.isinf(valor)):
                row_cells[i].text = "N/D"
            else:
                row_cells[i].text = formatear_valor(valor, es_dinero)

def tabla_proyectos_transpuesta(doc, dict_dfs, lista_ids, config_columnas, cols_agrupacion, col_id='Codproyecto'):
    df_filtrado, config_interna = _consolidar_bases(dict_dfs, lista_ids, config_columnas, cols_agrupacion, col_id)
    
    if df_filtrado.empty:
        doc.add_paragraph("No se encontraron datos para mostrar en la tabla.")
        return

    def calcular_grupo(grupo):
        resultado = {}
        for nombre_visual, (col_excel, es_dinero, metodo, col_peso) in config_interna.items():
            if metodo == 'calc':
                continue
                
            if col_excel in cols_agrupacion:
                if len(cols_agrupacion) == 1:
                    resultado[nombre_visual] = grupo.name
                else:
                    idx = cols_agrupacion.index(col_excel)
                    resultado[nombre_visual] = grupo.name[idx]
                continue
                
            special_methods = ['sum_last12_sells']
            if col_excel not in grupo.columns and col_excel is not None and metodo not in special_methods:
                resultado[nombre_visual] = np.nan
                continue
                
            if metodo == 'first':
                resultado[nombre_visual] = grupo[col_excel].iloc[0]
            elif metodo == 'sum':
                resultado[nombre_visual] = grupo[col_excel].sum(skipna=True)
            elif metodo == 'sum_last12_sells':
                if '__' in str(col_excel):
                    prefijo_base = str(col_excel).split('__')[0] + '__'
                else:
                    prefijo_base = ''
                    
                cols_coincidentes = []
                for c in grupo.columns:
                    c_str = str(c)
                    c_clean = c_str.split('__')[1] if '__' in c_str else c_str
                    c_clean_lower = c_clean.lower()
                    
                    if re.search(r'\d{4}-\d{2}-\d{2}', c_clean_lower):
                        if any(kw in c_clean_lower for kw in ['vend', 'vent']):
                            if not any(exc in c_clean_lower for exc in ['precio', 'metro', 'saldo', 'acum', 'total', 'promedio', 'vabr']):
                                fecha = pd.to_datetime(re.search(r'\d{4}-\d{2}-\d{2}', c_clean_lower).group(), errors='coerce')
                                if pd.notna(fecha):
                                    cols_coincidentes.append((c_str, fecha))
                                
                cols_coincidentes.sort(key=lambda x: x[1])
                ultimos_12 = [c[0] for c in cols_coincidentes[-12:]]
                
                if ultimos_12:
                    resultado[nombre_visual] = grupo[ultimos_12].sum(numeric_only=True).sum(skipna=True)
                else:
                    resultado[nombre_visual] = np.nan
            elif metodo == 'mean':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                # Si es dinero o área, el 0 es un dato faltante (NaN)
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.mean(skipna=True)
                
            elif metodo == 'median':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.median(skipna=True)
                
            elif metodo == 'max':
                resultado[nombre_visual] = pd.to_numeric(grupo[col_excel], errors='coerce').max(skipna=True)
                
            elif metodo == 'min':
                vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                    vals = vals.replace(0, np.nan)
                resultado[nombre_visual] = vals.min(skipna=True)
                
            elif metodo == 'avg12':
                resultado[nombre_visual] = grupo[col_excel].sum(skipna=True) / 12
                
            elif metodo in ['weighted_mean', 'weighted_average']:
                if col_peso and col_peso in grupo.columns:
                    df_valido = grupo.copy()
                    
                    # 1. Convertir a numérico
                    df_valido[col_excel] = pd.to_numeric(df_valido[col_excel], errors='coerce')
                    df_valido[col_peso] = pd.to_numeric(df_valido[col_peso], errors='coerce')
                    
                    # 2. Convertir ceros a NaN si es dinero o área
                    if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower():
                        df_valido[col_excel] = df_valido[col_excel].replace(0, np.nan)
                    
                    # 3. Eliminar cualquier fila que ahora sea NaN
                    df_valido = df_valido.dropna(subset=[col_excel, col_peso])
                    
                    valores_num = df_valido[col_excel]
                    pesos_num = df_valido[col_peso]
                    
                    # 4. Ignorar filas cuyo "peso" (ej. Unidades Disponibles) sea 0 o negativo
                    mascara_valida = (pesos_num > 0)
                    
                    if mascara_valida.sum() > 0:
                        resultado[nombre_visual] = np.average(valores_num[mascara_valida], weights=pesos_num[mascara_valida])
                    else:
                        resultado[nombre_visual] = np.nan
                else:
                    vals = pd.to_numeric(grupo[col_excel], errors='coerce')
                    if es_dinero or 'area' in str(col_excel).lower() or 'área' in str(col_excel).lower(): 
                        vals = vals.replace(0, np.nan)
                    resultado[nombre_visual] = vals.mean(skipna=True)
            elif metodo == 'concat':
                unique_values = grupo[col_excel].dropna().unique()
                if len(unique_values) > 1:
                    resultado[nombre_visual] = ', '.join(map(str, unique_values))
                elif len(unique_values) == 1:
                    resultado[nombre_visual] = unique_values[0]
                else:
                    resultado[nombre_visual] = np.nan
            else:
                resultado[nombre_visual] = grupo[col_excel].iloc[0]
                
        for nombre_visual, (col_excel, es_dinero, metodo, operacion) in config_interna.items():
            if metodo == 'calc':
                try:
                    resultado[nombre_visual] = operacion(resultado)
                except (ZeroDivisionError, KeyError, TypeError):
                    resultado[nombre_visual] = np.nan      
        return pd.Series(resultado)
    
    try:
        df_agrupado = df_filtrado.groupby(cols_agrupacion).apply(calcular_grupo, include_groups=False)
    except TypeError:
        df_agrupado = df_filtrado.groupby(cols_agrupacion).apply(calcular_grupo)
        
    df_agrupado = df_agrupado.reset_index(drop=True)
    
    num_rows = len(config_columnas)
    num_cols = len(df_agrupado) + 1 
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    nombres_columnas_keys = list(config_columnas.keys())
    
    for row_idx, nombre_visual in enumerate(nombres_columnas_keys):
        row_cells = table.rows[row_idx].cells
        
        p = row_cells[0].paragraphs[0]
        p.add_run(nombre_visual).bold = True

        for col_idx, (_, row_data) in enumerate(df_agrupado.iterrows(), start=1):
            valor = row_data.get(nombre_visual, np.nan)
            es_dinero = config_columnas[nombre_visual][2] 
            
            if pd.isna(valor) or (isinstance(valor, (int, float)) and np.isinf(valor)):
                row_cells[col_idx].text = "N/D"
            else:
                row_cells[col_idx].text = formatear_valor(valor, es_dinero)

def cargar_bases(archivo_csv, db_gi, db_cu, db_gi_comercial):
    if not os.path.exists(archivo_csv):
        raise FileNotFoundError(f"Error: No se encontró el archivo {archivo_csv}")
    if not os.path.exists(db_gi):
        raise FileNotFoundError(f"Error: No se encontró el archivo {db_gi}")
    if not os.path.exists(db_cu):
        raise FileNotFoundError(f"Error: No se encontró el archivo {db_cu}")
    if not os.path.exists(db_gi_comercial):
        raise FileNotFoundError(f"Error: No se encontró el archivo {db_gi_comercial}")

    print("Cargando Bases...")
    df_catastral = pd.read_csv(archivo_csv, low_memory=False, dtype={'LotCodigo': str})
    if 'LotCodigo' in df_catastral.columns:
        df_catastral['LotCodigo'] = df_catastral['LotCodigo'].astype(str).str.strip()
    if 'chip' in df_catastral.columns:
        df_catastral['chip'] = df_catastral['chip'].astype(str).str.strip()
    engine_gi = _select_excel_engine(db_gi)
    engine_gi_comercial = _select_excel_engine(db_gi_comercial)

    df_proyectos = pd.read_excel(db_gi, sheet_name='Base Proyectos', engine=engine_gi) if engine_gi else pd.read_excel(db_gi, sheet_name='Base Proyectos')
    df_inmuebles = pd.read_excel(db_gi, sheet_name='Base Formulada', engine=engine_gi) if engine_gi else pd.read_excel(db_gi, sheet_name='Base Formulada')
    df_etapas = pd.read_excel(db_gi, sheet_name='Base Etapas', engine=engine_gi) if engine_gi else pd.read_excel(db_gi, sheet_name='Base Etapas')
    df_amenidades = pd.read_excel(db_gi, sheet_name='Serv. Com.', engine=engine_gi) if engine_gi else pd.read_excel(db_gi, sheet_name='Serv. Com.')
    df_comercio_proyectos = pd.read_excel(db_gi_comercial, sheet_name='Base Proyecto', engine=engine_gi_comercial) if engine_gi_comercial else pd.read_excel(db_gi_comercial, sheet_name='Base Proyecto')
    df_comercio_inmuebles = pd.read_excel(db_gi_comercial, sheet_name='Base Inmuebles', engine=engine_gi_comercial) if engine_gi_comercial else pd.read_excel(db_gi_comercial, sheet_name='Base Inmuebles')
    df_comercio_proyectos = separar_coordenadas(df=df_comercio_proyectos, col_origen='Coordenadas Reales', col_lat='Latitud', col_lon='Longitud')

    df_cu_general = pd.read_excel(db_cu, sheet_name='Base General', skiprows=6)
    df_cu_acabados = pd.read_excel(db_cu, sheet_name='Base Acabados', skiprows=6)

    # 1. Galería Inmobiliaria (Vivienda)
    _rename_latest_regex(df_proyectos, r'^\s*P(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Pactual')
    _rename_latest_regex(df_inmuebles, r'^\s*P(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Pactual')
    _rename_latest_regex(df_inmuebles, r'^\s*D(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Dactual')
    _rename_latest_regex(df_inmuebles, r'^\s*V(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Vactual')
    _rename_latest_regex(df_etapas, r'^\s*EO(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'EOactual')

    # 2. Galería Inmobiliaria (Comercio)
    _rename_latest_regex(df_comercio_inmuebles, r'^\s*P(?:v|V)?(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Precio comercio actual')
    _rename_latest_regex(df_comercio_inmuebles, r'^\s*\$m2\s*(?:v|V)?(ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)(\d{2})\s*$', 'Precio m2 comercio actual')

    # 3. Coordenada Urbana
    _rename_latest_cu(df_cu_general, 'Precio', 'Precio Actual')
    _rename_latest_cu(df_cu_general, 'Precio metro cuadrado', 'Precio metro cuadrado Actual')
    _rename_latest_cu(df_cu_general, 'Vendidas', 'Vendidas Actual')
    _rename_latest_cu(df_cu_general, 'Saldo', 'Saldo Actual')

    _rename_coordinate_columns(df_proyectos)
    _rename_coordinate_columns(df_comercio_proyectos)
    _rename_coordinate_columns(df_cu_general)

    # Cast CU Coordinates to numeric safely
    df_cu_general['Latitud'] = pd.to_numeric(df_cu_general['Latitud'], errors='coerce')
    df_cu_general['Longitud'] = pd.to_numeric(df_cu_general['Longitud'], errors='coerce')
    df_cu_general = df_cu_general.dropna(subset=['Latitud', 'Longitud'])

    for col in df_cu_general.columns:
        if 'precio' in str(col).lower():
            df_cu_general[col] = pd.to_numeric(df_cu_general[col], errors='coerce') * 1000
    if '% Cuota Inicial' in df_proyectos.columns:
        df_proyectos['% Cuota Inicial'] = pd.to_numeric(df_proyectos['% Cuota Inicial'], errors='coerce') * 100
    if 'Fecha Inicio' in df_proyectos.columns:
        df_proyectos['Fecha Inicio'] = pd.to_datetime(pd.to_numeric(df_proyectos['Fecha Inicio'], errors='coerce'), origin='1899-12-30', unit='D').dt.strftime('%d/%m/%Y')
    if 'Fecha Terminacion Obra' in df_proyectos.columns:
        df_proyectos['Fecha Terminacion Obra'] = pd.to_datetime(pd.to_numeric(df_proyectos['Fecha Terminacion Obra'], errors='coerce'), origin='1899-12-30', unit='D').dt.strftime('%d/%m/%Y')
    df_inmuebles = crear_precio_por_m2(df_inmuebles)

    print("Bases cargadas. Procesando datos...")

    dict_dfs = {
        'Base Proyectos': df_proyectos,
        'Base Formulada': df_inmuebles,
        'Base Etapas': df_etapas,
        'Serv. Com.': df_amenidades,
        'Comercio_proyectos': df_comercio_proyectos,
        'Comercio_inmuebles': df_comercio_inmuebles,
        'Base General cu': df_cu_general,
        'Base Acabados cu': df_cu_acabados
    }
    return df_catastral, dict_dfs

def crear_word(
    df_catastral, dict_dfs, lotcodigo, nombre_predio, columna_id='LotCodigo', coordenadas=None, 
    radio_gi=1.5, radio_cu=1, radio_comercio=1.0, 
    excluir_codigos_gi_proyectos=None, excluir_codigos_cu=None, excluir_codigos_gi_comercio=None
):

    #Supuestos
    smmlv = 1750905  # Salario Mínimo Mensual Legal Vigente en Colombia para 2026
    if columna_id not in df_catastral.columns:
        print(f"Error: la columna '{columna_id}' no existe en la base catastral.")
        return

    lotcodigo_str = str(lotcodigo).strip()
    registro = df_catastral[df_catastral[columna_id].astype(str).str.strip().str.lower() == lotcodigo_str.lower()]
    if registro.empty:
        print(f"No se encontró ningún registro con el LotCodigo: {lotcodigo}")
        return
    
    # Creación documento
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    doc.add_heading(f'Estudio de mercado {nombre_predio}', 0)
    
    # --- SECCIÓN 1 ---
    doc.add_heading("1 Localización y características del predio", level=1)
    mapeo_columnas = {
        'direccion': ('Dirección', False),
        'chip': ('CHIP', False),
        'LotCodigo': ('Código de Lote', False),
        'area_terreno_m2': ('Área Terreno (m²)', False),
        'valor_m2_terreno': ('Valor m² Terreno', True),
        'valor_m2_referencia': ('Valor m² Referencia', True),
        'avaluo_catastral': ('Avalúo Catastral', True)
    }
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Valor'
    
    for col_csv, (nombre_visual, es_dinero) in mapeo_columnas.items():
        if col_csv in registro.columns or col_csv in ['LotCodigo', 'chip', 'avaluo_catastral']:
            if col_csv == 'LotCodigo':
                valor_final = str(lotcodigo)
            elif col_csv == 'chip':
                if 'chip' in registro.columns:
                    chips = registro['chip'].dropna().astype(str).unique().tolist()
                    valor_final = ', '.join(chips) if len(chips) > 0 else 'N/D'
                else:
                    valor_final = 'N/D'
            elif col_csv in ['avaluo_catastral', 'area_terreno_m2']:
                if col_csv in registro.columns:
                    suma_valores = pd.to_numeric(registro[col_csv], errors='coerce').sum(skipna=True)
                    valor_final = formatear_valor(suma_valores, es_dinero)
                else:
                    valor_final = 'N/D'
            else:
                try:
                    valor_original = registro[col_csv].dropna().iloc[0]
                except Exception:
                    valor_original = registro[col_csv].values[0] if col_csv in registro.columns and len(registro[col_csv].values) > 0 else None
                valor_final = formatear_valor(valor_original, es_dinero)

            row_cells = table.add_row().cells
            row_cells[0].text = nombre_visual
            row_cells[1].text = valor_final
    doc.add_paragraph("Fuente: Base catastral 2026. Unidad Administrativa Especial de Catastro Distrital.")
    doc.add_paragraph("Ilustración 1. Delimitación geográfica del predio.")
    generar_mapa_lote(doc, lotcodigo)
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo, Mayo 2026.")

    doc.add_heading("1.1 Manzana Catastral", level=2)
    doc.add_paragraph("Ilustración 2. Mediana del valor comercial por m² de terreno por manzana")
    generar_mapa_valor_referencia(doc, lotcodigo[:9])
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base en datos catastrales, Mayo 2026." \
    "https://serviciosgis.catastrobogota.gov.co/arcgis/rest/services/catastro/valorreferencia/MapServer/0")

    # --- SECCIÓN 2 ---
    doc.add_heading("2 Caracterización de la zona", level=1)
    doc.add_heading("2.1 Características", level=2)
    doc.add_heading("2.2 Altura", level=2)
    doc.add_paragraph("Ilustración 3. Número de pisos ponderado por manzana ")
    generar_mapa_altura_media(doc, lotcodigo[:9])
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base en datos catastrales, Mayo 2026." \
    "https://serviciosgis.catastrobogota.gov.co/arcgis/rest/services/catastro/alturamedia/MapServer/0")
    
    # --- SECCIÓN 3 ---
    doc.add_heading("3 Análisis de mercado inmobiliario vivienda", level=1)
    doc.add_heading("3.1 Metodología", level=2)
    doc.add_paragraph("Para el análisis del mercado inmobiliario de esta sección, se toma como fuente la base de proyectos de vivienda para RenoBo de las plataformas Galería Inmobiliaria y Coordenada Urbana, con fecha de corte a abril de 2026. Se realizó una identificación de proyectos activos a diferentes distancias, delimitando varios radios de influencia con intervalos de 500 m hasta llegar a un radio máximo de 2500 m. La ubicación de estos proyectos y la cantidad de proyectos encontrados por intervalos de distancia se encuentra a continuación:")
    doc.add_paragraph("Ilustración 4. Área analizada en el sector usando Galería Inmobiliaria")
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base de proyectos de Galería Inmobiliaria, Mayo 2026.")
    doc.add_paragraph("Ilustración 5. Área analizada en el sector usando Coordenada Urbana")
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base de proyectos de Coordenada Urbana, Mayo 2026.")
    doc.add_paragraph("Tabla 2. Cantidad proyectos encontrados por rango")

    doc.add_heading("3.2 Mercado inmobiliario.", level=2)
    doc.add_paragraph(f"A partir de identificar los proyectos activos más cercanos y que son competencia directa en ambas plataformas, " \
    f"se define delimitar el estudio a una circunferencia de {radio_cu} km para coordenada urbana y {radio_gi} km para Galería inmobiliaria sección vivienda alrededor del predio.  " \
    f"A continuación, se muestran algunas de las características principales de los proyectos competencia identificados en  galería inmobiliaria:")

    df_proyectos = dict_dfs['Base Proyectos']
    df_cu_general = dict_dfs['Base General cu']
    df_comercio_proyectos = dict_dfs['Comercio_proyectos']

    proyectos_cerca_gi = obtener_proyectos_cercanos(df_proyectos, coordenadas, radio_km=radio_gi, filtros=[('Activo', '==', 'Si')], fila_cod="Codproyecto", excluir_codigos=excluir_codigos_gi_proyectos)
    proyectos_cerca_cu = obtener_proyectos_cercanos(df_cu_general, coordenadas, radio_km=radio_cu, filtros=[('Segmento', '!=', 'SIN ASIGNAR'), ('Segmento', '!=', 'No Residencial'), ('Precio Actual', '!=', 0)], fila_cod="Id Proyecto", excluir_codigos=excluir_codigos_cu)
    comercio_cercano_gi = obtener_proyectos_cercanos(df_comercio_proyectos, coordenadas, radio_km=radio_comercio, filtros=[('Sub Zona', '!=', 'Soacha')], excluir_codigos=excluir_codigos_gi_comercio)
    comercio_cercano_cu = obtener_proyectos_cercanos(df_cu_general, coordenadas, radio_km=radio_comercio, fila_cod="Id Proyecto", filtros=[('usos', '==', 'No residencial')])
    
    configuracion_fichas = {
        "Nombre del Proyecto" : ('Base Proyectos', 'Proyecto', False, 'first', None),
        "Código del Proyecto": ('Base Proyectos', 'Codproyecto', False, 'first', None),
        "Barrio" : ('Base Proyectos', 'Barrio', False, 'first', None),
        "Dirección": ('Base Proyectos', 'Dirección', False, 'first', None),
        "Estrato": ('Base Proyectos', 'Estrato', False, 'first', None),
        "Constructor": ('Base Proyectos', 'Construye', False, 'first', None),
        "Financiera" : ('Base Proyectos', 'Financiera', False, 'first', None),
        "Fiduciaria" : ('Base Proyectos', 'Fiduciaria', False, 'first', None),
        "Descripción del Proyecto" : ('Serv. Com.', 'Descripcion del Proyecto', False, 'first', None),
        "Parqueaderos" : ('Serv. Com.', 'Parqueaderos', False, 'first', None),
        "Precio promedio Parqueaderos" : ('Base Formulada', 'Valor Garaje', False, 'weighted_mean', 'Oferta Total Inm.'),
        "Tot. Un. Proyecto" : ('Base Proyectos', 'Tot. Un. Proyecto', False, 'sum', None),
        "Área promedio (m²)" : ('Base Formulada', 'Area', False, 'weighted_mean', 'Oferta Total Inm.'),
        "Precio promedio ($)" : ('Base Formulada', 'Pactual', False, 'weighted_mean', 'Oferta Total Inm.'),
        "Amenidades" : ('Serv. Com.', 'Otro', False, 'first', None),
        "Otros Usos" : ('Base Proyectos', 'Otro Uso', False, 'first', None),
        "Entrega" : ('Base Proyectos', 'Entrega (Obra Gris o Acabados)', False, 'first', None),
        "Incremento de precio entre piso" : ('Base Proyectos', 'Incremento de precio entre piso', False, 'first', None),
        "% Cuota Inicial" : ('Base Proyectos', '% Cuota Inicial', False, 'first', None),
        "Tipo de Vivienda" : ('Base Proyectos', 'Tipo de Vivienda ( Vis o No Vis )', False, 'first', None),
        "Estado Actual" : ('Base Proyectos', 'Estado', False, 'first', None),
        "Teléfono" : ('Base Proyectos', 'Tel.', False, 'first', None),
        "Fecha de inicio" : ('Base Proyectos', 'Fecha Inicio', False, 'first', None),
        "Fecha de entrega" : ('Base Proyectos', 'Fecha Terminacion Obra', False, 'first', None)
    }

    cols_agrupacion = ['Proyecto']
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_gi, configuracion_fichas, cols_agrupacion)
    doc.add_paragraph("A continuación se muestran las comparaciones de los proyectos competencia de coordenada urbana")
    configuracion_fichas_cu = {
        "Nombre del Proyecto" : ('Base General cu', 'Nombre Del Proyecto', False, 'first', None),
        "Id Proyecto": ('Base General cu', 'Id Proyecto', False, 'first', None),
        "Barrio" : ('Base General cu', 'Barrio', False, 'first', None),
        "Dirección": ('Base General cu', 'Dirección', False, 'first', None),
        "Tipo" : ('Base General cu', 'Segmento', False, 'concat', None),
        "Estrato": ('Base General cu', 'Estrato', False, 'first', None),
        "Constructor": ('Base General cu', 'Nombre Del Constructor', False, 'first', None),
        "Financiera" : ('Base General cu', 'Entidad Financiera (Cred. Constructor)', False, 'first', None),
        "Fiduciaria" : ('Base General cu', 'Fiduciaria', False, 'first', None),
        "Tot. Un. Proyecto" : ('Base General cu', 'Unidades Por Tipo', False, 'sum', None),
        "Área promedio (m²)" : ('Base General cu', 'Área Por Tipo', False, 'weighted_mean', 'Unidades Por Tipo'),
        "Precio promedio ($)" : ('Base General cu', 'Precio Actual', False, 'weighted_mean', 'Unidades Por Tipo'),
        "Entrega" : ('Base General cu', 'condicion_entrega', False, 'concat', None),
        "Estado Actual" : ('Base General cu', 'Estado', False, 'first', None),
        "Teléfono" : ('Base General cu', 'Teléfono', False, 'first', None),
        "Fecha de inicio Ventas" : ('Base General cu', 'Fecha Inicio Ventas', False, 'first', None),
        "Fecha de entrega" : ('Base General cu', 'Fecha De Entrega', False, 'first', None)
    }

    cols_agrupacion = ['Nombre Del Proyecto']
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_cu, configuracion_fichas_cu, cols_agrupacion, col_id='Id Proyecto')
    
    # Subsecciones de 3.2
    doc.add_heading("3.2.1 Evolución del Mercado", level=3)
    doc.add_heading("3.2.1.1 Evolución de ventas y oferta", level=4)
    
    doc.add_paragraph("A continuación se muestra la evolución histórica de ventas mensuales (Barras) y la oferta disponible (Línea) de los últimos 12 meses para proyectos cercanos de Galería Inmobiliaria.")
    doc.add_paragraph("Ilustración 6. Evolución histórica de ventas y oferta de proyectos cercanos de Galería Inmobiliaria")
    # Gráfico Doble Eje - Galería Inmobiliaria
    grafico_historico_doble_eje(
        doc=doc,
        dict_dfs=dict_dfs,
        lista_ids=proyectos_cerca_gi,
        cols_agrupacion=['Tipo VIS'],
        nombre_base='Base Formulada',
        col_id='Codproyecto',
        prefijo_barras='V',
        formato_barras='prefijo_mes',
        prefijo_linea='D',
        formato_linea='prefijo_mes',
        etiqueta_y_barras='Ventas Mensuales (unidades)',
        etiqueta_y_linea='Oferta Disponible (unidades)',
        titulo_barras='Ventas Totales',
        titulo_linea='Oferta Disponible',
        titulo_grafico='Histórico de Oferta y Ventas (Galería Inmobiliaria)'
    )
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base en datos de Galería Inmobiliaria, Mayo 2026.")
    doc.add_paragraph("Tabla 6. Resumen comercial por proyectos de Galería Inmobiliaria")
    tabla_comercial_tipo = { 
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Ventas acumuladas últimos 12 meses': ('Base Formulada', 'V.Acum.', False, 'sum', None),
        'Ventas promedio últimos 12 meses': ('Base Formulada', 'V.Acum.', False, 'calc', lambda res: res['Ventas acumuladas últimos 12 meses'] / 12 if res.get('Ventas acumuladas últimos 12 meses', 0) > 0 else np.nan),
        'Oferta total en inmuebles': ('Base Formulada', 'Oferta Total Inm.', False, 'sum', None),
        'Oferta disponible': ('Base Formulada', 'Dactual', False, 'sum', None),
        'Disponibilidad porcentual (%)': ('Base Formulada', 'Disponibilidad porcentual (%)', False, 'calc', lambda res: (res['Oferta disponible'] / res['Oferta total en inmuebles'] * 100) if res.get('Oferta total en inmuebles', 0) > 0 else np.nan),
    } 
    cols_agrupacion = ['Tipo VIS'] 
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_gi, tabla_comercial_tipo, cols_agrupacion)
    doc.add_paragraph("Tabla 7. Distribución del estado actual de la oferta de Galería inmobiliaria")
    tabla_estado_actual = [
        ("Estado Oferta","Base Etapas","EOactual")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_gi, tabla_estado_actual)
    doc.add_paragraph("A continuación se muestra la evolución histórica de ventas mensuales (Barras) y la oferta disponible (Línea) de los últimos 12 meses para proyectos cercanos de Coordenada Urbana.")
    doc.add_paragraph("Ilustración 7. Evolución histórica de ventas y oferta de proyectos cercanos de Coordenada Urbana")
    # Gráfico Doble Eje - Coordenada Urbana
    grafico_historico_doble_eje(
        doc=doc,
        dict_dfs=dict_dfs,
        lista_ids=proyectos_cerca_cu,
        cols_agrupacion=['Segmento'],
        nombre_base='Base General cu',
        col_id='Id Proyecto',
        prefijo_barras='Vendidas',
        formato_barras='texto_fecha',
        prefijo_linea='Saldo',
        formato_linea='texto_fecha',
        etiqueta_y_barras='Ventas Mensuales (unidades)',
        etiqueta_y_linea='Oferta Disponible (unidades)',
        titulo_barras='Ventas Totales',
        titulo_linea='Oferta Disponible',
        titulo_grafico='Histórico de Oferta y Ventas (Coordenada Urbana)'
    )
    doc.add_paragraph("Fuente: Elaboración Propia RenoBo con base en datos de Coordenada Urbana, Mayo 2026.")
    doc.add_paragraph("Tabla 8. Resumen comercial por proyectos de Coordenada Urbana")
    tabla_comercial_tipo_cu = {
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Ventas acumuladas ult. 12 meses (und)': ('Base General cu', 'Vendidas Actual', False, 'sum_last12_sells', None),
        'Ventas promedio ult. 12 meses (und)': ('Base General cu', 'Vendidas Actual', False, 'calc', lambda res: res['Ventas acumuladas ult. 12 meses (und)'] / 12 if res.get('Ventas acumuladas ult. 12 meses (und)', 0) > 0 else np.nan),
        'Oferta total (und)': ('Base General cu', 'Unidades Por Tipo', False, 'sum', None),
        'Oferta disponible (und)': ('Base General cu', 'Saldo Actual', False, 'sum', None),
        'Disponibilidad porcentual (%)': ('Base General cu', 'Disponibilidad porcentual (%)', False, 'calc', lambda res: (res['Oferta disponible (und)'] / res['Oferta total (und)'] * 100) if res.get('Oferta total (und)', 0) > 0 else np.nan),
    }
    cols_agrupacion = ['Segmento']
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_cu, tabla_comercial_tipo_cu, cols_agrupacion, col_id='Id Proyecto')
    
    doc.add_paragraph("Tabla 9.Distribución del estado actual de la oferta de Coordenada Urbana")
    tabla_estado_actual_cu = [
        ("Estado Oferta","Base General cu","Estado")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_cu, tabla_estado_actual_cu, col_id='Id Proyecto')
    doc.add_heading("3.2.1.2 Evolución de precio m²", level=4)
    doc.add_paragraph("A continuación se muestra la evolución del precio promedio por m² de los proyectos cercanos de Galería inmobiliaria.")
    doc.add_paragraph("Ilustración 8. Evolución histórica del precio promedio por m² de proyectos cercanos de Galería Inmobiliaria")
    grafico_evolucion_temporal(
        doc=doc,
        dict_dfs=dict_dfs,
        lista_ids=proyectos_cerca_gi,
        cols_agrupacion=['Tipo VIS'],
        col_id='Codproyecto',
        tipo_columnas='custom',
        prefijo_columnas='Pm',
        formato_columnas='prefijo_mes',
        etiqueta_y='Precio promedio por m² ($)',
        incluir_tabla=False
    )
    doc.add_paragraph("A continuación se muestra la evolución del precio promedio por m² de los proyectos cercanos de Coordenada Urbana.")
    doc.add_paragraph("Ilustración 9. Evolución histórica del precio promedio por m² de proyectos cercanos de Coordenada Urbana")
    grafico_evolucion_temporal(
        doc=doc,
        dict_dfs=dict_dfs,
        lista_ids=proyectos_cerca_cu,
        cols_agrupacion=['Segmento'],
        col_id='Id Proyecto',
        nombre_base='Base General cu',
        tipo_columnas='custom',
        prefijo_columnas='Precio metro cuadrado',
        formato_columnas='texto_fecha',
        etiqueta_y='Precio promedio por m² ($)',
        incluir_tabla=False
    )
    
    doc.add_heading("3.2.2 Información comercial por proyecto", level=3)
    doc.add_paragraph("En esta sección se muestra información acerca de las unidades de ventas de los últimos 12 meses por cada proyecto y cada tipo. Además, muestra la oferta total y disponible por proyecto")
    doc.add_paragraph("Tabla 10. Resumen comercial por proyectos de Galería Inmobiliaria")
    tabla_comercial_proyecto = {
        'Nombre del Proyecto': ('Base Formulada', 'Proyecto', False, 'first', None), 
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Ventas acumuladas últimos 12 meses': ('Base Formulada', 'V.Acum.', False, 'sum', None),
        'Ventas promedio últimos 12 meses': ('Base Formulada', 'V.Acum.', False, 'calc', lambda res: res['Ventas acumuladas últimos 12 meses'] / 12),
        'Oferta total (und)': ('Base Formulada', 'Oferta Total Inm.', False, 'sum', None),
        'Oferta disponible (und)': ('Base Formulada', 'Dactual', False, 'sum', None)
    }

    cols_agrupacion = ['Codproyecto', 'Tipo VIS'] 
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_gi, tabla_comercial_proyecto, cols_agrupacion)

    doc.add_paragraph("Tabla 11. Resumen comercial por proyectos de Coordenada Urbana")
    tabla_comercial_proyecto_cu = {
        'Nombre del Proyecto': ('Base General cu', 'Nombre Del Proyecto', False, 'first', None), 
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Ventas acumuladas ult. 12 meses (und)': ('Base General cu', 'Vendidas Actual', False, 'sum_last12_sells', None),
        'Ventas promedio ult. 12 meses (und)': ('Base General cu', 'Vendidas Actual', False, 'calc', lambda res: res['Ventas acumuladas ult. 12 meses (und)'] / 12),
        'Oferta total (und)': ('Base General cu', 'Unidades Por Tipo', False, 'sum', None),
        'Oferta disponible (und)': ('Base General cu', 'Saldo Actual', False, 'sum', None),
        'Disponibilidad porcentual (%)': ('Base General cu', 'Disponibilidad porcentual (%)', False, 'calc', lambda res: (res['Oferta disponible (und)'] / res['Oferta total (und)'] * 100) if res.get('Oferta total (und)', 0) > 0 else np.nan)
    }

    cols_agrupacion = ['Id Proyecto', 'Segmento'] 
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_cu, tabla_comercial_proyecto_cu, cols_agrupacion, col_id='Id Proyecto')

    doc.add_heading("3.2.3 Información precios por proyecto", level=3)
    doc.add_paragraph("Tabla 12. Precios por proyecto encontrados usando Galería Inmobiliaria")
    tabla_precios_proyecto = {
        'Proyecto': ('Base Formulada', 'Proyecto', False, 'first', None),
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Área (m²)': ('Base Formulada', 'Area', False, 'weighted_average', 'Oferta Total Inm.'),
        'Precio promedio ($)': ('Base Formulada', 'Pactual', True, 'weighted_mean', 'Oferta Total Inm.'),
        'Precio promedio SMMLV': ('Base Formulada', 'Pactual', False, 'calc', lambda res: res['Precio promedio ($)'] / smmlv if res.get('Precio promedio ($)', 0) > 0 else np.nan),
        'Precio promedio ($/m²)': ('Base Formulada', 'Pactual', True, 'calc', lambda res: (res['Precio promedio ($)'] / res['Área (m²)']) if res.get('Área (m²)', 0) > 0 else np.nan),
        'Año de inicio': ('Base Proyectos', 'Fecha Inicio', False, 'first', None),
        'Año de entrega': ('Base Proyectos', 'Fecha Terminacion Obra', False, 'first', None)
    }
    cols_agrupacion = ['Codproyecto', 'Tipo VIS'] 
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_gi, tabla_precios_proyecto, cols_agrupacion)
    doc.add_paragraph("Tabla 13. Precios por proyecto encontrados usando Coordenada Urbana")
    
    tabla_proyectos_cu_1 = {
        'Proyecto': ('Base General cu', 'Nombre Del Proyecto', False, 'first', None),
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Área promedio (m²)': ('Base General cu', 'Área Por Tipo', False, 'weighted_average', 'Unidades Por Tipo'),
        'Precio promedio ($)': ('Base General cu', 'Precio Actual', True, 'weighted_mean', 'Unidades Por Tipo'),
        'Precio promedio SMMLV': ('Base General cu', 'Precio Actual', False, 'calc', lambda res: res['Precio promedio ($)'] / smmlv if res.get('Precio promedio ($)', 0) > 0 else np.nan),
        'Precio promedio ($/m²)': ('Base General cu', 'Precio metro cuadrado Actual', True, 'calc', lambda res: (res['Precio promedio ($)'] / res['Área promedio (m²)']) if res.get('Área promedio (m²)', 0) > 0 else np.nan),
        'Fecha de inicio Ventas' : ('Base General cu', 'Fecha Inicio Ventas', False, 'first', None),
        'Fecha de entrega' : ('Base General cu', 'Fecha De Entrega', False, 'first', None)
    }
    cols_agrupacion = ['Id Proyecto', 'Segmento']
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_cu, tabla_proyectos_cu_1, cols_agrupacion, col_id='Id Proyecto')

    doc.add_paragraph("Tabla 14. Resumen de precios en galería inmobiliaria")
    tabla_precios_tipo = {
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Área promedio (m²)': ('Base Formulada', 'Area', False, 'weighted_average', 'Oferta Total Inm.'),
        'Área mediana (m²)': ('Base Formulada', 'Area', False, 'median', None),
        'Área máxima (m²)': ('Base Formulada', 'Area', False, 'max', None),
        'Área mínima (m²)': ('Base Formulada', 'Area', False, 'min', None),
        'Precio promedio ($)': ('Base Formulada', 'Pactual', True, 'weighted_mean', 'Oferta Total Inm.'),
        'Precio máximo ($)': ('Base Formulada', 'Pactual', True, 'max', None),
        'Precio mínimo ($)': ('Base Formulada', 'Pactual', True, 'min', None),
        'Precio promedio SMMLV': ('Base Formulada', 'Pactual', False, 'calc', lambda res: res['Precio promedio ($)'] / smmlv if res.get('Precio promedio ($)', 0) > 0 else np.nan),
        'Precio máximo SMMLV': ('Base Formulada', 'Pactual', False, 'calc', lambda res: res['Precio máximo ($)'] / smmlv if res.get('Precio máximo ($)', 0) > 0 else np.nan),
        'Precio mínimo SMMLV': ('Base Formulada', 'Pactual', False, 'calc', lambda res: res['Precio mínimo ($)'] / smmlv if res.get('Precio mínimo ($)', 0) > 0 else np.nan),
        'Precio promedio ($/m²)': ('Base Formulada', '$m2', True, 'weighted_mean', 'Oferta Total Inm.'),
        'Precio máximo ($/m²)': ('Base Formulada', '$m2', True, 'max', None),
        'Precio mínimo ($/m²)': ('Base Formulada', '$m2', True, 'min', None),
    }
    cols_agrupacion = ['Tipo VIS'] 
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_gi, tabla_precios_tipo, cols_agrupacion)
    doc.add_paragraph("Tabla 15. Resumen de precios en coordenada urbana")
    tabla_proyectos_cu_3 = {
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Área promedio (m²)': ('Base General cu', 'Área Por Tipo', False, 'weighted_average', 'Unidades Por Tipo'),
        'Área mediana (m²)': ('Base General cu', 'Área Por Tipo', False, 'median', None),
        'Área máxima (m²)': ('Base General cu', 'Área Por Tipo', False, 'max', None),
        'Área mínima (m²)': ('Base General cu', 'Área Por Tipo', False, 'min', None),
        'Precio promedio ($)': ('Base General cu', 'Precio Actual', True, 'weighted_mean', 'Unidades Por Tipo'),
        'Precio máximo ($)': ('Base General cu', 'Precio Actual', True, 'max', None),
        'Precio mínimo ($)': ('Base General cu', 'Precio Actual', True, 'min', None),
        'Precio promedio SMMLV': ('Base General cu', 'Precio Actual', False, 'calc', lambda res: res['Precio promedio ($)'] / smmlv if res.get('Precio promedio ($)', 0) > 0 else np.nan),
        'Precio máximo SMMLV': ('Base General cu', 'Precio Actual', False, 'calc', lambda res: res['Precio máximo ($)'] / smmlv if res.get('Precio máximo ($)', 0) > 0 else np.nan),
        'Precio mínimo SMMLV': ('Base General cu', 'Precio Actual', False, 'calc', lambda res: res['Precio mínimo ($)'] / smmlv if res.get('Precio mínimo ($)', 0) > 0 else np.nan),
        'precio promedio ($/m²)': ('Base General cu', 'Precio metro cuadrado Actual', True, 'weighted_mean', 'Área Por Tipo'),
        'precio máximo ($/m²)': ('Base General cu', 'Precio metro cuadrado Actual', True, 'max', None),
        'precio mínimo ($/m²)': ('Base General cu', 'Precio metro cuadrado Actual', True, 'min', None),
    }
    cols_agrupacion = ['Segmento'] 
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_cu, tabla_proyectos_cu_3, cols_agrupacion, col_id='Id Proyecto')
    
    doc.add_heading("3.2.5 Características de los inmuebles", level=3)
    doc.add_paragraph("Tabla 16. Características de los inmuebles por proyecto encontrados usando Galería Inmobiliaria")
    tabla_caracteristicas_proyecto = {
        'Proyecto': ('Base Proyectos', 'Proyecto', False, 'first', None),
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Area promedio (m²)': ('Base Formulada', 'Area', False, 'weighted_average', 'Oferta Total Inm.'),
        'Número de habitaciones': ('Base Formulada', 'Alc.', False, 'weighted_average', 'Oferta Total Inm.'),
        'Número de baños completos': ('Base Formulada', 'Baños Completos', False, 'weighted_average', 'Oferta Total Inm.'),
        'Entrega': ('Base Proyectos', 'Entrega (Obra Gris o Acabados)', False, 'first', None),
        'Parqueaderos': ('Serv. Com.', 'Parqueaderos', False, 'first', None),
    }
    cols_agrupacion = ['Codproyecto', 'Tipo VIS']
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_gi, tabla_caracteristicas_proyecto, cols_agrupacion)
    doc.add_paragraph("Tabla 17. Características de los inmuebles por proyecto encontrados usando Coordenada Urbana")
    tabla_proyectos_cu_2 = {
        'Proyecto': ('Base General cu', 'Nombre Del Proyecto', False, 'first', None),
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Área promedio (m²)': ('Base General cu', 'Área Por Tipo', False, 'weighted_average', 'Unidades Por Tipo'),
        'Alcobas': ('Base General cu', 'Alcobas', False, 'weighted_average', 'Unidades Por Tipo'),
        'Baños': ('Base General cu', 'Baños', False, 'weighted_average', 'Unidades Por Tipo'),    
        'Tipo de Acabado': ('Base General cu', 'condicion_entrega', False, 'first', None),
    }
    cols_agrupacion = ['Id Proyecto', 'Segmento']
    tabla_proyectos(doc, dict_dfs, proyectos_cerca_cu, tabla_proyectos_cu_2, cols_agrupacion, col_id='Id Proyecto')

    doc.add_paragraph("Tabla 18. Resumen de características en galería inmobiliaria")
    tabla_caracteristicas_tipo = {
        'Tipo': ('Base Formulada', 'Tipo VIS', False, 'first', None),
        'Área promedio (m²)': ('Base Formulada', 'Area', False, 'weighted_average', 'Oferta Total Inm.'),
        'Promedio de habitaciones': ('Base Formulada', 'Alc.', False, 'weighted_average', 'Oferta Total Inm.'),
        #'Número de habitaciones máximo': ('Base Formulada', 'Alc.', False, 'max', None),
        #'Número de habitaciones mínimo': ('Base Formulada', 'Alc.', False, 'min', None),
        'Promedio de Baños Completos': ('Base Formulada', 'Baños Completos', False, 'weighted_average', 'Oferta Total Inm.'),
        #'Número de baños completos máximo': ('Base Formulada', 'Baños Completos', False, 'max', None),
        #'Número de baños completos mínimo': ('Base Formulada', 'Baños Completos', False, 'min', None),
    }
    cols_agrupacion = ['Tipo VIS']
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_gi, tabla_caracteristicas_tipo, cols_agrupacion) 

    doc.add_paragraph("Tabla 19. Resumen de características en coordenada urbana")
    tabla_proyectos_cu_4 = {
        'Tipo': ('Base General cu', 'Segmento', False, 'first', None),
        'Área promedio (m²)': ('Base General cu', 'Área Por Tipo', False, 'weighted_average', 'Unidades Por Tipo'),
        'Alcobas': ('Base General cu', 'Alcobas', False, 'weighted_average', 'Unidades Por Tipo'),
        'Baños': ('Base General cu', 'Baños', False, 'weighted_average', 'Unidades Por Tipo'),   
    }
    cols_agrupacion = ['Segmento']
    tabla_proyectos_transpuesta(doc, dict_dfs, proyectos_cerca_cu, tabla_proyectos_cu_4, cols_agrupacion, col_id='Id Proyecto')

    doc.add_heading("3.2.6 Especificaciones y Acabados Interiores de los inmuebles", level=3)
    doc.add_paragraph("A continuación se muestran las especificaciones y acabados interiores de los proyectos competencia de Galería Inmobiliaria y Coordenada Urbana que se entregan con semi acabados o acabados.")
    doc.add_paragraph("Tabla 20. Especificaciones y Acabados Interiores de proyectos competencia de Galería Inmobiliaria")
    proyectos_cerca_gi = obtener_proyectos_cercanos(df_proyectos, coordenadas, radio_km=radio_gi, filtros=[('Entrega (Obra Gris o Acabados)', '!=', 'Obra Gris')])
    config_acabados = [
        ("Cocina Integral", "Serv. Com.", "Coc. Integral"),
        ("Mueble de Cocina", "Serv. Com.", "Mueble Cocina"),
        ("Mesón de Cocina", "Serv. Com.", "Mesón Cocina"),
        ("Muros de Cocina", "Serv. Com.", "Muros Cocina"),
        ("Estufa", "Serv. Com.", "Estufa Gas / Eléc."),
        ("Horno", "Serv. Com.", "Horno Gas / Eléc."),
        ("Extractor", "Serv. Com.", "Extractor"),
        ("Calentador", "Serv. Com.", "Cal. Gas / Eléc."),
        ("Lavadero", "Serv. Com.", "Lavad."),
        ("Instalación Lavadora/Secadora", "Serv. Com.", "Inst. Lavad. / Secadora"),
        ("Mueble Lavamanos", "Serv. Com.", "Mueble Lavamanos"),
        ("Mesón de Baños", "Serv. Com.", "Mesón Baños"),
        ("Muros de Baños", "Serv. Com.", "Muros Baños"),
        ("Piso de Baños", "Serv. Com.", "Piso Baño"),
        ("Pared de Ducha", "Serv. Com.", "Pared Ducha"),
        ("Tinas", "Serv. Com.", "Tinas"),
        ("Sanitario Alcoba Principal", "Serv. Com.", "Tipo De Sanitario Alcoba Princ."),
        ("Sanitario Otros Baños", "Serv. Com.", "Tipo Sanitario Otros Baños"),
        ("Tipo Lavamanos", "Serv. Com.", "Tipo Lavamanos"),
        ("Grifería Lavamanos", "Serv. Com.", "Tipo Griferia Lav."),
        ("Grifería Ducha", "Serv. Com.", "Tipo Griferia Ducha"),
        ("Tipo Ducha", "Serv. Com.", "Tipo Ducha"),
        ("Piso Zona Social", "Serv. Com.", "Piso Zona Social"),
        ("Piso Alcobas", "Serv. Com.", "Piso Alcoba"),
        ("Piso Halles Vivienda", "Serv. Com.", "Piso Halles Viv."),
        ("Muros Interiores", "Serv. Com.", "Muros Interiores"),
        ("Ventanería", "Serv. Com.", "Vent."),
        ("Puertas", "Serv. Com.", "Pta"),
        ("Carpintería Puertas/Closets", "Serv. Com.", "Carpinteria Puertas/Closets"),
        ("Chimeneas", "Serv. Com.", "Chimeneas"),
        ("Aire Acondicionado", "Serv. Com.", "Aire Acond.")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_gi, config_acabados)

    doc.add_paragraph("Tabla 21. Especificaciones y Acabados Interiores de los proyectos competencia de Coordenada Urbana")
    proyectos_cerca_cu = obtener_proyectos_cercanos(df_cu_general, coordenadas, radio_km=radio_cu, filtros=[('condicion_entrega', '==', 'Acabados')], fila_cod="Id Proyecto")
    config_acabados_interiores = [
        ("Pisos Alcobas", "Base Acabados cu", "Pisos Alcobas"),
        ("Muros Alcobas", "Base Acabados cu", "Muros Alcobas"),
        ("Cielo Raso Alcobas", "Base Acabados cu", "Cielo Raso Alcobas"),
        ("Pisos Cocina", "Base Acabados cu", "Pisos Cocina"),
        ("Muros Cocina", "Base Acabados cu", "Muros Cocina"),
        ("Cielo Raso Cocina", "Base Acabados cu", "Cielo Raso Cocina"),
        ("Tipo de Cocina", "Base Acabados cu", "Tipo Cocina"),
        ("Mesón Cocina", "Base Acabados cu", "Meson Cocina"),
        ("Muebles Cocina", "Base Acabados cu", "Muebles Cocina"),
        ("Dotación Cocina", "Base Acabados cu", "Dotación Cocina"),
        ("Pisos Baños", "Base Acabados cu", "Pisos Baños"),
        ("Muros Baños", "Base Acabados cu", "Muros Baños"),
        ("Cielo Raso Baños", "Base Acabados cu", "Cielo Raso Baños"),
        ("Mesón Baños", "Base Acabados cu", "Meson Baños"),
        ("División Baño", "Base Acabados cu", "Division Baño"),
        ("Lavamanos", "Base Acabados cu", "Lavamanos"),
        ("Mueble Lavamanos", "Base Acabados cu", "Mueble Lavamanos"),
        ("Lavadero", "Base Acabados cu", "Lavadero"),
        ("Puerta Principal", "Base Acabados cu", "Puerta Principal"),
        ("Puertas Interiores", "Base Acabados cu", "Puertas Interiores"),
        ("Material de Clósets", "Base Acabados cu", "Closet Material"),
        ("Cornisa", "Base Acabados cu", "Cornisa")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_cu, config_acabados_interiores, col_id='Id Proyecto')

    doc.add_heading("3.2.7 Amenidades de los proyectos", level=2)
    doc.add_paragraph("En esta sección se muestran las amenidades que ofrecen los proyectos competencia de Galería Inmobiliaria y Coordenada Urbana. Para esto, se muestran tablas con la proporción de proyectos que ofrecen cada amenidad.")
    doc.add_paragraph("Tabla 22. Amenidades de proyectos encontrados en galería inmobiliaria")
    proyectos_cerca_gi = obtener_proyectos_cercanos(df_proyectos, coordenadas, radio_km=radio_gi)
    config_amenidades = [
        ("Portería", "Serv. Com.", "Portería"),
        ("Salón Social", "Serv. Com.", "Salón Social"),
        ("Parque Infantil", "Serv. Com.", "Parque Inf."),
        ("Canchas", "Serv. Com.", "Canchas"),
        ("Gimnasio", "Serv. Com.", "Gimnasio"),
        ("Ofrece Piscina", "Serv. Com.", "Ofrece Piscina"),
        ("Zonas Húmedas", "Serv. Com.", "Zonas Humedas"),
        ("Zona de BBQ", "Serv. Com.", "BBQ"),
        ("Zona de Mascotas", "Serv. Com.", "Zona Mascotas"),
        ("Parqueaderos", "Serv. Com.", "Parqueaderos"),
        ("Depósito", "Serv. Com.", "Depósito"),
        ("Ascensores por Torre", "Serv. Com.", "No. Ascensores x Torre Minimo"),
        ("Planta Eléctrica", "Serv. Com.", "Planta Eléct."),
        ("Cuarto/Shut de Basuras", "Serv. Com.", "Shut Basuras"),
        ("Tipo de Urbanización", "Serv. Com.", "Tipo Urbanizacion"),
        ("Sistema Estructural", "Serv. Com.", "Tipo Est."),
        ("Fachada", "Serv. Com.", "Fachada"),
        ("Pisos en Halles Comunales", "Serv. Com.", "Piso Halles Com.")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_gi, config_amenidades)

    doc.add_paragraph("Tabla 23. Amenidades de proyectos encontrados en Coordenada Urbana")
    proyectos_cerca_cu = obtener_proyectos_cercanos(df_cu_general, coordenadas, radio_km=radio_cu, fila_cod="Id Proyecto", filtros= None)
    config_zonas_dummies = [
        ("BBQ", "Base Acabados cu", "Zona BBQ", True),
        ("Estación de Carga", "Base Acabados cu", "Zona Estación de Carga", True),
        ("Balcón", "Base Acabados cu", "Zona Balcón", True),
        ("Centro Cultural", "Base Acabados cu", "Zona Centro Cultural", True),
        ("Circulaciones", "Base Acabados cu", "Zona Circulaciones", True),
        ("Vestier / Casilleros", "Base Acabados cu", "Zona Vestier/Casilleros", True),
        ("Verde Picnic", "Base Acabados cu", "Zona Verde Picnic", True),
        ("Sede Social", "Base Acabados cu", "Zona Sede Social", True),
        ("Ascensores", "Base Acabados cu", "Zona Ascensores", True),
        ("Gimnasio", "Base Acabados cu", "Zona Gimnasio", True),
        ("Guardería", "Base Acabados cu", "Zonas Guardería", True),
        ("Cicloruta", "Base Acabados cu", "Zona Cicloruta", True),
        ("Piscina", "Base Acabados cu", "Zona Piscina", True),
        ("Bar", "Base Acabados cu", "Zona Bar", True),
        ("Canchas Múltiples", "Base Acabados cu", "Zona Canchas Multiples", True),
        ("Cancha Squash", "Base Acabados cu", "Zona Cancha Squash", True),
        ("Cancha de Fútbol", "Base Acabados cu", "Zona Cancha de Futbol", True),
        ("Lobby Recepción", "Base Acabados cu", "Zona Lobby Recepción", True),
        ("Oficina Administración", "Base Acabados cu", "Zona Oficina Administración", True),
        ("Espejo de Agua", "Base Acabados cu", "Zona Espejo de Agua", True),
        ("Parque Infantil", "Base Acabados cu", "Zonas Parque Infantil", True),
        ("Citofonía", "Base Acabados cu", "Zonas Citofona", True),
        ("Aire Acondicionado", "Base Acabados cu", "Zonas Aire Acondicionado", True),
        ("Cinema", "Base Acabados cu", "Zona Cinema", True),
        ("Sendero de Trote", "Base Acabados cu", "Zona Sendero de Trote", True),
        ("Terraza Sky Club", "Base Acabados cu", "Zona Terraza Sky Club", True),
        ("Fuente", "Base Acabados cu", "Zona Fuente", True),
        ("Red Contra Incendios", "Base Acabados cu", "Zona Red Contra Incendios", True),
        ("Zona de Pesca", "Base Acabados cu", "Zona Zona de Pesca", True),
        ("Kiosco", "Base Acabados cu", "Zona Kiosco", True),
        ("Cuarto de Basuras", "Base Acabados cu", "Zona Cuarto de Basuras", True),
        ("Cancha Tenis", "Base Acabados cu", "Zona Cancha Tenis", True),
        ("Shut de Basuras", "Base Acabados cu", "Zonas_Shut_Basuras", True),
        ("Tanque Abastecimiento de Agua", "Base Acabados cu", "Zona Tanque Abastecimiento de Agua", True),
        ("Salón Comunal", "Base Acabados cu", "Zona Salón Comunal", True),
        ("Caldera", "Base Acabados cu", "Zona Caldera", True),
        ("Golfito", "Base Acabados cu", "Zona Golfito", True),
        ("Portería", "Base Acabados cu", "Zona Portería", True),
        ("Salón de Juegos", "Base Acabados cu", "Zona Salón de Juegos", True),
        ("Campo de Golf", "Base Acabados cu", "Zona_Campo_de_Golf", True),
        ("Lavandería", "Base Acabados cu", "Zona Lavandería", True),
        ("Bicicletero", "Base Acabados cu", "Zona Bicicletero", True),
        ("Zonas Húmedas", "Base Acabados cu", "Zonas Húmedas", True),
        ("Salón Especiales / Sala de Junta", "Base Acabados cu", "Zona Salón Especiales Sala de Junta", True),
        ("Planta Eléctrica de Emergencia", "Base Acabados cu", "Zona Planta Electrica de Emergencia", True),
        ("Plazoleta", "Base Acabados cu", "Zona Plazoleta", True),
        ("Centro Computo", "Base Acabados cu", "Zona Centro Computo", True),
        ("Cafetería", "Base Acabados cu", "Zona Cafetería", True),
        ("Restaurante", "Base Acabados cu", "Zona Restaurante", True),
        ("Auditorios", "Base Acabados cu", "Zona Auditorios", True),
        ("Capilla", "Base Acabados cu", "Zona Capilla", True),
        ("Batería de Baños", "Base Acabados cu", "Zona Batería de Baños", True),
        ("Biblioteca", "Base Acabados cu", "Zona Biblioteca", True),
        ("Cuarto Técnico / de Gases", "Base Acabados cu", "Zona cuarto Técnico/de Gases", True),
        ("Mascotas", "Base Acabados cu", "Zona Mascotas", True),
        ("Cuarto de Equipos", "Base Acabados cu", "Zona Cuarto de Equipos", True),
        ("Niños", "Base Acabados cu", "Zona de Niños", True),
        ("Área Lounge", "Base Acabados cu", "Zona Área Lounge", True),
        ("Cocina Social", "Base Acabados cu", "Zona Cocina Social", True),
        ("Cancha Voleibol", "Base Acabados cu", "Zona Cancha Voleibol", True),
        ("Fogata", "Base Acabados cu", "Zona Fogata", True),
        ("Tanque Reserva", "Base Acabados cu", "Zona Tanque Reserva", True),
        ("Taller Bici", "Base Acabados cu", "Zona Taller Bici", True),
        ("Roof Garden", "Base Acabados cu", "Zona Roof Garden", True),
        ("Salas de Estudio", "Base Acabados cu", "Zona Salas de Estudio", True),
        ("Escaleras", "Base Acabados cu", "Zona Escaleras", True),
        ("Mini-market", "Base Acabados cu", "Zona Mini-market", True),
        ("Conjunto Cerrado", "Base Acabados cu", "Zona Conjunto Cerrado", True),
        ("Sala Múltiple", "Base Acabados cu", "Zona Sala Múltiple", True),
        ("Teatrino", "Base Acabados cu", "Zona Teatrino", True),
        ("Co-working", "Base Acabados cu", "Zona Co-working", True),
        ("Yoga", "Base Acabados cu", "Zona Yoga", True),
        ("Spa", "Base Acabados cu", "Zona Spa", True),
        ("Sauna", "Base Acabados cu", "Zona Sauna", True),
        ("Depósito", "Base Acabados cu", "Zona Déposito", True),
        ("Teatro", "Base Acabados cu", "Zona Teatro", True),
        ("Vías de Acceso", "Base Acabados cu", "Zona vías de acceso", True),
        ("Jacuzzi", "Base Acabados cu", "Zona Jacuzzi", True),
        ("N/A", "Base Acabados cu", "Zona N/A", True),
        ("Muro de Escalar", "Base Acabados cu", "Zona Muro escalar", True),
        ("Monta Coches", "Base Acabados cu", "Zona Monta Coches", True),
        ("Enfermería", "Base Acabados cu", "Zona Enfermería", True),
        ("Despacho Peatonales", "Base Acabados cu", "Zona Despacho peatonales", True),
        ("Pista Bicicross", "Base Acabados cu", "Zona Pista Bicicross", True)
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_cu, config_zonas_dummies, col_id='Id Proyecto')

    doc.add_heading("3.2.8 Especificaciones Constructivas de los proyectos", level=4)
    doc.add_paragraph("Tabla 24. Características de proyectos de competencia de Galería Inmobiliaria")
    config_constructivas_gi = [
        ("Tipo de Urbanización", "Serv. Com.", "Tipo Urbanizacion"),
        ("Sistema Estructural", "Serv. Com.", "Tipo Est."),
        ("Fachada", "Serv. Com.", "Fachada"),
        ("Pisos en Halles Comunales", "Serv. Com.", "Piso Halles Com.")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_gi, config_constructivas_gi)
    doc.add_paragraph("Tabla 25. Características de proyectos de competencia de Coordenada Urbana")
    
    config_constructivas_cu = [
        ("Placa Entre Piso", "Base Acabados cu", "Placa Entre Piso"),
        ("Sistema Constructivo", "Base Acabados cu", "Sistema Constructivo"),
        ("Subestructura Cimentación", "Base Acabados cu", "Subestructura Cimentacion"), 
        ("Subestructura Sótanos", "Base Acabados cu", "Subestructura Sotanos"),
        ("Divisiones Interiores", "Base Acabados cu", "Divisiones Interiores"),
        ("Fachada", "Base Acabados cu", "Fachada"),
        ("Cubiertas", "Base Acabados cu", "Cubiertas"),
        ("Ventanería", "Base Acabados cu", "Ventanería")
    ]
    tabla_proporciones(doc, dict_dfs, proyectos_cerca_cu, config_constructivas_cu, col_id='Id Proyecto')

    doc.add_heading("3.3 Mercado inmobiliario no residencial", level=2)
    doc.add_paragraph(f"En esta sección se analizan los proyectos con usos no residenciales que se encuentran cerca al predio, los cuales pueden ser competencia directa o indirecta dependiendo del caso. Este análisis se realizó trazando una circunferencia de {radio_comercio} metros alrededor del predio")
    doc.add_paragraph("Tabla 26. Proyectos con usos no residenciales competencia de Galería Inmobiliaria")

    tabla_comercial_proyecto_gi = {
        'Nombre del Proyecto': ('Comercio_inmuebles', 'Proyecto', False, 'first', None),
        'Tipo': ('Comercio_inmuebles', 'Tipo', False, 'first', None),
        'Oferta total (m²)': ('Comercio_inmuebles', 'Oferta Lanzada Inm. M2', False, 'sum', None),
        'Oferta disponible (m²)': ('Comercio_inmuebles', 'D m2 Inmueble', False, 'sum', None),
        '% Disponible': ('Comercio_inmuebles', None, False, 'calc', lambda res: (res['Oferta disponible (m²)'] / res['Oferta total (m²)'] * 100) if res.get('Oferta total (m²)', 0) > 0 else np.nan),
        'Fecha inicio de ventas': ('Comercio_inmuebles', 'Fecha Lanzamiento', False, 'first', None),
        'Precio promedio ($/m²)': ('Comercio_inmuebles', 'Precio m2 comercio actual', True, 'weighted_mean', 'Total Oferta Base Inm.'),
        'Precio promedio ($)': ('Comercio_inmuebles', 'Precio comercio actual', True, 'weighted_mean', 'Total Oferta Base Inm.')
    }
    cols_agrupacion = ['Codproyecto', 'Tipo']
    tabla_proyectos(doc, dict_dfs, comercio_cercano_gi, tabla_comercial_proyecto_gi, cols_agrupacion)
    doc.add_paragraph("Tabla 27. Proyectos con usos no residenciales competencia de Coordenada Urbana")
    tabla_comercial_proyecto_cu = {
        'Nombre del Proyecto': ('Base General cu', 'Nombre Del Proyecto', False, 'first', None),
        'Tipología': ('Base General cu', 'Segmento', False, 'first', None),
        'Tipo Por Etapa': ('Base General cu', 'Tipo Por Etapa', False, 'first', None),
        'Oferta total (m²)': ('Base General cu', 'Área Por Tipo', False, 'sum', None),
        'Unidades disponibles': ('Base General cu', 'Saldo Actual', False, 'sum', None),
        'Oferta disponible (m²)': ('Base General cu', None, False, 'calc', lambda res: (res['Área Por Tipo'] * res['Unidades disponibles']) if res.get('Área Por Tipo') is not None and res.get('Unidades disponibles') is not None else np.nan),
        '% Disponible': ('Base General cu', None, False, 'calc', lambda res: (res['Oferta disponible (m²)'] / res['Oferta total (m²)'] * 100) if res.get('Oferta total (m²)', 0) > 0 else np.nan),
        'Fecha inicio de ventas': ('Base General cu', 'Fecha Inicio Ventas', False, 'first', None),
        'Precio promedio ($/m²)': ('Base General cu', 'Precio metro cuadrado Actual', True, 'weighted_mean', 'Área Por Tipo'),
        'Precio promedio ($)': ('Base General cu', 'Precio Actual', True, 'weighted_mean', 'Área Por Tipo')
    }
    cols_agrupacion = ['Id Proyecto', 'Segmento','Tipo Por Etapa']
    tabla_proyectos(doc, dict_dfs, comercio_cercano_cu, tabla_comercial_proyecto_cu, cols_agrupacion, col_id='Id Proyecto')
    
    # --- SECCIÓN 4 ---
    doc.add_heading("4 Conclusiones", level=1)
    #Falta Incluír Conclusiones y recomendaciones.
    os.makedirs(nombre_predio, exist_ok=True)
    nombre_archivo = f"{nombre_predio}/Reporte_{nombre_predio}.docx"
    doc.save(nombre_archivo)
    print(f"Documento '{nombre_archivo}' generado con éxito.")